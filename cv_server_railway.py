#!/usr/bin/env python3
"""
cv_server_railway.py  —  v2.3-groq
LLM: Groq (primario) → Gemini (fallback) → Claude (fallback)

Formulario multi-pantalla:
  1a. Email only → detecta si existe
  2a. Si existe → "¡Hola de nuevo!" + botones Buscar ahora / Mañana 9am
  1.  Si nuevo → formulario completo + botón Buscar ahora
"""

import os
import io
import re
import logging
import requests
from datetime import datetime, timezone
from typing import NamedTuple
from flask import Flask, request, jsonify, render_template_string

# Google Drive / OAuth
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# Claude (calidad — CV y textos que van a empresas)
import anthropic

# DOCX
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Búsqueda de ofertas reales (Fase 3.0)
from real_jobs import buscar_ofertas_reales

# ─────────────────────────────────────────────
# CONFIGURACIÓN — solo variables de entorno
# ─────────────────────────────────────────────

# ── LLM: Groq (primario) ──────────────────────
GROQ_API_KEY = os.environ["GROQ_API_KEY"]                          # requerido
GROQ_MODEL   = os.getenv("GROQ_MODEL", "openai/gpt-oss-120b")      # llama-3.3-70b-versatile lo retiro Groq el 16-ago-2026

# ── LLM: Gemini (fallback opcional) ──────────
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
# gemini-1.5-flash y gemini-2.0-flash dan 404 desde antes del 28-ago-2026.
# Google responde "no longer available, please update to models/gemini-3.6-flash".
GEMINI_MODEL   = os.getenv("GEMINI_MODEL", "gemini-3.6-flash")

# ── LLM: Claude (fallback opcional) ──────────
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY", "")
CLAUDE_MODEL   = os.getenv("CLAUDE_MODEL", "claude-haiku-4-5")  # fallback si Groq falla.
# OJO: el default era "claude-3-haiku-20240307", RETIRADO el 19-abr-2026. La cadena de
# fallback devolvia 404 y nadie lo habia notado porque casi nunca se ejercita.

# ── Claude para el CV (calidad — va a empresas; Groq queda de fallback) ──
# Haiku 4.5: barato (~$0,02/CV) y sigue bien el prompt de adaptación.
CV_MODEL = os.getenv("CV_MODEL", "claude-haiku-4-5")
# Carta de presentación: Sonnet 4.6 (mejor prosa, ~$0,04/carta). Va a empresas.
CARTA_MODEL = os.getenv("CARTA_MODEL", "claude-sonnet-4-6")

# ── Google Drive ──────────────────────────────
GOOGLE_CLIENT_ID     = os.environ["GOOGLE_CLIENT_ID"]
GOOGLE_CLIENT_SECRET = os.environ["GOOGLE_CLIENT_SECRET"]
GOOGLE_REFRESH_TOKEN = os.environ["GOOGLE_REFRESH_TOKEN"]
# Dos carpetas SEPARADAS a proposito: los CV Master son la fuente de verdad y
# se leen; los CV generados son salida y se acumulan a varios por dia. Si
# comparten carpeta, en semanas es imposible distinguirlos y se corre el riesgo
# de leer como master un CV generado.
FOLDER_CV_MASTERS    = os.getenv("FOLDER_CV_MASTERS", "1duJA_G3lLbOqiUYoSJcsXAvbtJUdcmzR")
FOLDER_CV_GENERADOS  = os.getenv("FOLDER_CV_GENERADOS", "1tHuVOIz3ratjRp8AmHsF0kGVpmy9DocY")

# ── Notion ────────────────────────────────────
NOTION_TOKEN = os.environ["NOTION_TOKEN"]
NOTION_DB_USUARIOS = os.getenv("NOTION_DB_USUARIOS", "")
NOTION_DB_OFERTAS  = os.getenv("NOTION_DB_OFERTAS", "33d11515-f4b2-8176-947b-000bbafd1ca7")

# ── Webhooks n8n ──────────────────────────────
# OJO (28ago2026): los paths `buscar-ahora` y `nuevo-usuario` NO EXISTEN en la
# instancia. Se barrieron los 10 workflows nodo a nodo. Lo que la documentacion
# llamaba WF1 nunca llego a estar dado de alta, y estas llamadas se comian un 404
# en silencio. El unico webhook vivo que lanza una busqueda para un usuario es
# `buscar-para-user`, dentro del workflow de PROD `CsvmtPcLVmGIZg6C`.
N8N_HOST = os.getenv("N8N_HOST", "https://n8n-asistente-correo.onrender.com")
WEBHOOK_BUSCAR_AHORA = os.getenv(
    "WEBHOOK_BUSCAR_AHORA", f"{N8N_HOST}/webhook/buscar-para-user"
)

# ─────────────────────────────────────────────
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)


class CVError(Exception):
    """Error de negocio del cv-server: status HTTP + mensaje.

    El core (generar_cv_core, ...) la lanza; cada capa HTTP (Flask/FastAPI)
    la mapea a su formato de respuesta. Ver docs/ADR-001."""

    def __init__(self, status: int, message: str):
        super().__init__(message)
        self.status = status
        self.message = message


# ══════════════════════════════════════════════
# CAPA LLM — Groq primario, Gemini/Claude fallback
# ══════════════════════════════════════════════

class RespuestaLLM(NamedTuple):
    """Respuesta de un LLM junto al modelo que la generó DE VERDAD.

    Sin esto no se puede saber si el CV lo escribió Claude o el fallback
    de Groq: los endpoints reportaban el modelo configurado, no el usado.
    """
    contenido: str
    modelo:    str


def call_llm(prompt: str) -> RespuestaLLM:
    """Llama a Groq; si falla intenta Gemini y luego Claude."""

    # ── 1. Groq ──────────────────────────────
    try:
        resp = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Content-Type":  "application/json",
                "Authorization": f"Bearer {GROQ_API_KEY}",
            },
            json={
                "model":      GROQ_MODEL,
                "messages":   [{"role": "user", "content": prompt}],
                "max_tokens": 4096,
                "temperature": 0.7,
            },
            timeout=30,
        )
        resp.raise_for_status()
        content = resp.json()["choices"][0]["message"]["content"]
        logger.info("LLM: Groq OK (%s)", GROQ_MODEL)
        return RespuestaLLM(content, GROQ_MODEL)
    except Exception as e:
        logger.warning("Groq falló: %s — probando fallbacks", e)

    # ── 2. Gemini (fallback) ──────────────────
    if GEMINI_API_KEY:
        try:
            resp = requests.post(
                f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent",
                params={"key": GEMINI_API_KEY},
                json={"contents": [{"parts": [{"text": prompt}]}]},
                timeout=30,
            )
            resp.raise_for_status()
            content = resp.json()["candidates"][0]["content"]["parts"][0]["text"]
            logger.info("LLM: Gemini fallback OK (%s)", GEMINI_MODEL)
            return RespuestaLLM(content, GEMINI_MODEL)
        except Exception as e:
            logger.warning("Gemini fallback falló: %s — probando Claude", e)

    # ── 3. Claude (fallback) ──────────────────
    if CLAUDE_API_KEY:
        try:
            resp = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type":      "application/json",
                    "x-api-key":         CLAUDE_API_KEY,
                    "anthropic-version": "2023-06-01",
                },
                json={
                    "model":      CLAUDE_MODEL,
                    "max_tokens": 4096,
                    "messages":   [{"role": "user", "content": prompt}],
                },
                timeout=30,
            )
            resp.raise_for_status()
            content = resp.json()["content"][0]["text"]
            logger.info("LLM: Claude fallback OK (%s)", CLAUDE_MODEL)
            return RespuestaLLM(content, CLAUDE_MODEL)
        except Exception as e:
            logger.error("Claude fallback falló: %s", e)

    raise RuntimeError("Todos los LLMs fallaron. Revisa las API keys y el estado de los servicios.")


# ── Capa CALIDAD: Claude primario para el CV (lo que va a empresas) ──
_anthropic_client = None

def get_anthropic_client():
    global _anthropic_client
    if _anthropic_client is None:
        if not CLAUDE_API_KEY:
            raise RuntimeError("CLAUDE_API_KEY no configurada")
        _anthropic_client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
    return _anthropic_client


def call_claude(prompt: str, model: str, max_tokens: int = 4096) -> str:
    """Llama a Claude vía SDK oficial. Para CV/carta donde la calidad importa."""
    resp = get_anthropic_client().messages.create(
        model=model,
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": prompt}],
    )
    return "".join(b.text for b in resp.content if b.type == "text")


def call_llm_calidad(prompt: str, model: str = CV_MODEL, max_tokens: int = 4096) -> RespuestaLLM:
    """Claude primario; si falla (rate limit, red o sin key) cae a Groq.
    Para el CV y textos que van a una empresa — mejor que Groq, ~$0,02/CV."""
    try:
        contenido = call_claude(prompt, model=model, max_tokens=max_tokens)
        logger.info("LLM calidad: Claude OK (%s)", model)
        return RespuestaLLM(contenido, model)
    except Exception as e:
        logger.warning("Claude falló (%s) — cayendo a Groq", e)
        return call_llm(prompt)


# ══════════════════════════════════════════════
# GOOGLE DRIVE
# ══════════════════════════════════════════════

def get_drive_service():
    creds = Credentials(
        token=None,
        refresh_token=GOOGLE_REFRESH_TOKEN,
        token_uri="https://oauth2.googleapis.com/token",
        client_id=GOOGLE_CLIENT_ID,
        client_secret=GOOGLE_CLIENT_SECRET,
        scopes=["https://www.googleapis.com/auth/drive"],
    )
    creds.refresh(Request())
    return build("drive", "v3", credentials=creds)


def subir_cv_a_drive(docx_bytes: bytes, nombre_archivo: str) -> str:
    service = get_drive_service()
    file_metadata = {
        "name":    nombre_archivo,
        "parents": [FOLDER_CV_GENERADOS],
    }
    media = MediaIoBaseUpload(
        io.BytesIO(docx_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    file = service.files().create(
        body=file_metadata, media_body=media, fields="id, webViewLink"
    ).execute()

    # Hacer público (solo lectura)
    service.permissions().create(
        fileId=file["id"],
        body={"role": "reader", "type": "anyone"},
    ).execute()

    return file.get("webViewLink", "")


# MimeTypes de Google Docs que necesitan export en vez de get_media
_GDOC_EXPORT = {
    "application/vnd.google-apps.document":       "text/plain",
    "application/vnd.google-apps.presentation":   "text/plain",
    "application/vnd.google-apps.spreadsheet":    "text/csv",
}


# ── Guardrail de veracidad: cifras inventadas ────
# El prompt YA prohibe inventar metricas y el modelo lo hizo igual (caso real:
# "serving millions of users", cifra que no existe en el Master). Una instruccion
# es una peticion, no una garantia: la salida se VERIFICA contra la fuente.

# Los guardrails viven en `guardrails.py`. Se reexportan aqui para no cambiar la
# superficie publica del modulo: los endpoints y los tests los siguen viendo en
# `cv_server_railway`.
from guardrails import (  # noqa: F401
    DESCRIPCION_MINIMA,
    construir_titular,
    detectar_cifras_no_respaldadas,
    detectar_experiencia_mal_atribuida,
    detectar_skills_no_respaldadas,
    detectar_tecnologias_no_respaldadas,
    detectar_titular_fuera_de_contrato,
    evaluar_descripcion_oferta,
    _TEC_ALIAS,
    _TEC_PATRONES,
    _tecnologias_en,
)

class MasterElegido(NamedTuple):
    """Master seleccionado para un idioma, con la URL de la que sale.

    Sin la url no se puede reportar CUAL master se usó: la respuesta acababa
    devolviendo siempre el master inglés aunque el CV fuese en español.
    """
    file_id: str
    url:     str


def elegir_master(usuario: dict, idioma: str) -> MasterElegido:
    """Elige la fuente del master segun el idioma. Pura: no toca Drive.
    idioma='en' -> 'CV Master URL' (ingles); cualquier otro -> 'CV Master URL ES'
    (con fallback al master ingles si no hay version española configurada)."""
    if idioma == "en":
        file_id = (usuario.get("cv_master_file_id") or "").strip()
        url = usuario.get("cv_master_url", "") or ""
    else:  # 'es' (o cualquier otro) -> master español, con fallback al ingles
        file_id = ""
        url = usuario.get("cv_master_url_es", "") or ""
        if not url:
            file_id = (usuario.get("cv_master_file_id") or "").strip()
            url = usuario.get("cv_master_url", "") or ""

    # Si no hay file_id directo, extraerlo de la URL (link de Drive/Docs)
    if not file_id and url:
        m = re.search(r'/d/([a-zA-Z0-9_-]+)', url) or re.search(r'id=([a-zA-Z0-9_-]+)', url)
        if m:
            file_id = m.group(1)

    return MasterElegido(file_id, url)


class MasterCV(NamedTuple):
    """Texto del master y la URL de la que se leyó de verdad."""
    texto: str
    url:   str


def leer_cv_master_desde_drive(usuario: dict, idioma: str = "es") -> MasterCV:
    """Descarga el CV master en texto plano desde Drive, eligiendo la fuente segun idioma."""
    service = get_drive_service()

    file_id, url = elegir_master(usuario, idioma)

    if not file_id:
        return MasterCV("", "")

    try:
        # Detectar mimeType para saber cómo extraer el texto
        file_meta = service.files().get(fileId=file_id, fields="mimeType, name", supportsAllDrives=True).execute()
        mime = file_meta.get("mimeType", "")
        name = file_meta.get("name", "")

        if mime in _GDOC_EXPORT:
            # Google Docs nativos → exportar a texto
            export_mime = _GDOC_EXPORT[mime]
            req = service.files().export_media(fileId=file_id, mimeType=export_mime)
        else:
            # Archivos binarios (DOCX, PDF, etc.) → get_media
            req = service.files().get_media(fileId=file_id)

        buf = io.BytesIO()
        from googleapiclient.http import MediaIoBaseDownload
        dl = MediaIoBaseDownload(buf, req)
        done = False
        while not done:
            _, done = dl.next_chunk()
        buf.seek(0)

        # DOCX es un ZIP, NO texto plano: hay que parsearlo con python-docx.
        # Decodificar sus bytes como utf-8 devuelve basura ("PK...word/document.xml")
        # y el LLM nunca ve la experiencia real → CV genérico.
        DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if mime == DOCX_MIME or name.lower().endswith(".docx"):
            doc = Document(buf)
            partes = [p.text for p in doc.paragraphs if p.text.strip()]
            # Las skills suelen ir en tablas → también hay que extraerlas
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        if celda.text.strip():
                            partes.append(celda.text)
            return MasterCV("\n".join(partes), url)

        # Texto plano u otros formatos legibles
        return MasterCV(buf.read().decode("utf-8", errors="replace"), url)
    except Exception as e:
        logger.warning("No se pudo leer CV master (file_id=%s): %s", file_id, e)
        return MasterCV("", url)


# ══════════════════════════════════════════════
# IDIOMA
# ══════════════════════════════════════════════

import re as _re_idioma

_ES_ACENTOS = _re_idioma.compile(r"[ñáéíóúü¿¡]", _re_idioma.IGNORECASE)
_ES_PALABRAS = {
    "experiencia", "equipo", "desarrollo", "empresa", "puesto", "requisitos",
    "conocimientos", "años", "trabajo", "ofrecemos", "buscamos", "gestión",
    "liderazgo", "desarrollador", "programador", "aplicaciones", "datos",
    "proyecto", "cliente", "habilidades", "capacidad", "valorable",
    "imprescindible", "nivel", "sector", "jornada", "remoto",
}
_EN_PALABRAS = {
    "experience", "team", "development", "company", "position", "requirements",
    "skills", "years", "work", "we", "you", "our", "developer", "engineer",
    "manage", "ability", "strong", "knowledge", "including", "required",
    "preferred", "remote", "build", "design", "role", "looking",
}


def _señales_idioma(texto: str) -> tuple:
    """Cuenta señales de español e inglés en un texto. Devuelve (es, en).
    Los acentos y signos ¿¡ pesan doble: son señal fuerte de español."""
    t = (texto or "").lower()
    if not t.strip():
        return (0, 0)
    palabras = set(_re_idioma.findall(r"[a-záéíóúñü]+", t))
    es = len(_ES_ACENTOS.findall(t)) * 2
    es += sum(1 for w in _ES_PALABRAS if w in palabras)
    en = sum(1 for w in _EN_PALABRAS if w in palabras)
    return (es, en)


def detectar_idioma(*textos) -> str:
    """Heuristica simple: devuelve 'es' o 'en' segun señales del texto de la oferta.
    Empate -> 'es' (mercado principal de la usuaria)."""
    es = en = 0
    for t in textos:
        e, n = _señales_idioma(t)
        es += e
        en += n
    return "en" if en > es else "es"


def idioma_de_oferta(puesto: str, descripcion: str, empresa: str) -> str:
    """Idioma del anuncio, priorizando el PUESTO.

    El titulo del puesto viene tal cual del anuncio, en su idioma. La descripcion,
    en cambio, la reescribe la tarea programada casi siempre en español, asi que
    ahogaba la señal del titulo (caso Revolut, 23jul2026: titulo ingles, carta en
    español). El puesto pesa x3, pero no es absoluto: una descripcion con señal
    española muy marcada todavia puede ganar (oferta española titulada en ingles).

    Un idioma explicito (body o campo Idioma de Notion) manda sobre esta deteccion:
    esta funcion es solo el ultimo recurso. Empate y vacio -> 'es'."""
    es_p, en_p = _señales_idioma(puesto)
    if es_p != en_p:
        # El titulo tiene señal neta: manda. No lo contamina la descripcion.
        return "en" if en_p > es_p else "es"
    # Titulo vacio o ambiguo: caemos a la descripcion y la empresa.
    es_d, en_d = _señales_idioma(descripcion)
    es_e, en_e = _señales_idioma(empresa)
    return "en" if (en_d + en_e) > (es_d + es_e) else "es"


def _slug(texto: str) -> str:
    """Slug en minúsculas sin acentos para nombres de archivo."""
    s = (texto or "").lower().strip()
    for a, b in (("á", "a"), ("é", "e"), ("í", "i"), ("ó", "o"),
                 ("ú", "u"), ("ñ", "n"), ("ü", "u")):
        s = s.replace(a, b)
    s = _re_idioma.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return s


def _nombre_archivo_cv(nombre: str, puesto: str) -> str:
    """Convención: cv-<nombre>-<puesto>-<año>.docx (ej: cv-veronica-serna-frontend-developer-2026.docx)."""
    partes = ["cv", _slug(nombre) or "candidato"]
    puesto_slug = _slug(puesto)
    if puesto_slug:
        partes.append(puesto_slug)
    partes.append(str(datetime.now(timezone.utc).year))
    return "-".join(partes) + ".docx"


def _tiene_algun_master(usuario: dict) -> bool:
    """True si el usuario tiene configurado un master en cualquier idioma."""
    return bool(
        (usuario.get("cv_master_file_id") or "").strip()
        or (usuario.get("cv_master_url") or "").strip()
        or (usuario.get("cv_master_url_es") or "").strip()
    )


# ══════════════════════════════════════════════
# NOTION
# ══════════════════════════════════════════════

def notion_headers():
    return {
        "Authorization":  f"Bearer {NOTION_TOKEN}",
        "Notion-Version": "2022-06-28",
        "Content-Type":   "application/json",
    }


# Campo de Notion (Users) con las direcciones ADICIONALES por las que entran ofertas
# del mismo usuario. `Email` sigue siendo la principal.
#
# El nombre lo escribe una persona en Notion, asi que se aceptan las variantes
# razonables (singular/plural, mayusculas). Exigir coincidencia exacta produce el
# peor fallo posible: el campo existe, el codigo no lo ve y NO hay error — el
# usuario simplemente se queda sin alias. Paso el 28jul2026 con "Email Alias".
CAMPO_EMAILS_ALIAS = os.getenv("CAMPO_EMAILS_ALIAS", "Emails alias")
_ALIAS_ACEPTADOS = {"emails alias", "email alias", "emails_alias", "email_alias",
                    "emails alternativos", "email alternativo"}


def _es_campo_de_alias(nombre: str) -> bool:
    return nombre.strip().lower() in _ALIAS_ACEPTADOS or nombre == CAMPO_EMAILS_ALIAS


def campos_alias_candidatos() -> list:
    """Nombres con los que INTENTAR la consulta a Notion, en orden.

    Tolerar el nombre al LEER no sirve de nada si al PREGUNTAR se usa uno solo:
    Notion filtra por nombre exacto y devuelve 400 si no existe. Paso el 28jul2026
    con 'Email alias' (la propiedad) frente a 'Emails alias' (lo que se consultaba):
    la busqueda por alias nunca encontraba nada y no habia error visible.
    """
    orden = [CAMPO_EMAILS_ALIAS, "Email alias", "Emails alias",
             "Emails alternativos", "Email alternativo"]
    vistos, out = set(), []
    for n in orden:
        if n.lower() not in vistos:
            vistos.add(n.lower()); out.append(n)
    return out

_SEPARADORES_EMAIL = re.compile(r"[,;\n\r]+")
_ES_EMAIL = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


def emails_de_usuario(props: dict) -> set:
    """Todas las direcciones por las que se puede identificar a un usuario.

    Una persona con dos buzones es UN usuario con dos emails, no dos usuarios: dos
    registros derivan (masters distintos, campos a medias) y el CV sale distinto
    segun por donde entre la oferta. Caso real, 28jul2026: ver ADR-003.
    """
    emails = set()
    principal = (props.get("Email", {}) or {}).get("email") or ""
    if principal.strip():
        emails.add(principal.strip().lower())

    crudo = ""
    for nombre, campo in props.items():
        if _es_campo_de_alias(nombre):
            crudo += "".join(t.get("plain_text", "") for t in ((campo or {}).get("rich_text") or []))
            crudo += ","
    for trozo in _SEPARADORES_EMAIL.split(crudo):
        t = trozo.strip().lower()
        if t and _ES_EMAIL.match(t):
            emails.add(t)
    return emails


def usuario_tiene_email(props: dict, email: str) -> bool:
    """¿Este usuario responde a esta direccion? Comparacion EXACTA.

    Importa que sea exacta: el filtro `contains` de Notion es de subcadena, asi que
    'vero@gmail.com' casa con 'notvero@gmail.com'. Sin esta verificacion final un
    usuario podria recibir el CV de otro.
    """
    return (email or "").strip().lower() in emails_de_usuario(props)


def _consultar_usuario(filtro: dict, email: str, verificar: bool) -> dict | None:
    """Una consulta a Users. Si verificar=True confirma el email exacto en Python."""
    try:
        resp = requests.post(
            f"https://api.notion.com/v1/databases/{NOTION_DB_USUARIOS}/query",
            headers=notion_headers(),
            json={"filter": filtro, "page_size": 5 if verificar else 1},
            timeout=15,
        )
    except requests.RequestException as e:
        logger.warning("Notion query fallo (%s): %s", filtro.get("property"), e)
        return None
    if resp.status_code != 200:
        # Un 400 aqui suele significar que el campo no existe todavia en la base.
        logger.warning("Notion query error %s (%s): %s",
                       resp.status_code, filtro.get("property"), resp.text[:200])
        return None
    for page in resp.json().get("results", []):
        if not verificar or usuario_atiende(page.get("properties", {}), email):
            return page
    return None


def usuario_atiende(props: dict, email: str) -> bool:
    """¿Este registro atiende esta direccion? Exige que este ACTIVO.

    Desactivar un registro tiene que significar "ya no atiende a nadie". Sin esto,
    un duplicado desactivado seguia ganando: la busqueda mira primero el `Email`
    exacto, lo encontraba ahi, y nunca llegaba al alias del registro bueno. Costo
    un CV generado con la identidad equivocada el 28jul2026 (ver ADR-003).

    Sin columna `Activo` se asume activo, para no romper bases que no la tengan.
    """
    activo = props.get("Activo", {}).get("checkbox", True) if "Activo" in props else True
    return bool(activo) and usuario_tiene_email(props, email)


def buscar_usuario_por_email(email: str) -> dict | None:
    """Consulta Notion por email (principal o alias). Devuelve el perfil o None."""
    if not NOTION_DB_USUARIOS:
        return None
    # verificar=True tambien en la pasada por Email exacto: hay que comprobar que el
    # registro este ACTIVO, cosa que el filtro de Notion no hace.
    page = _consultar_usuario(
        {"property": "Email", "email": {"equals": email}}, email, verificar=True
    )
    if page is None:
        # Segunda pasada por los alias. `contains` es de subcadena, de ahi el
        # verificar=True: se confirma la coincidencia exacta en Python.
        for campo in campos_alias_candidatos():
            page = _consultar_usuario(
                {"property": campo, "rich_text": {"contains": email}},
                email, verificar=True,
            )
            if page is not None:
                break
    if page is None:
        return None
    p = page.get("properties", {})
    return {
        "notion_id":          page.get("id", ""),
        "nombre":             (p.get("Name", {}).get("title") or [{}])[0].get("plain_text", ""),
        "email":              p.get("Email", {}).get("email", ""),
        "email_cv":           (p.get("Email CV", {}).get("email", "")
                               or (p.get("Email CV", {}).get("rich_text") or [{}])[0].get("plain_text", "")),
        "activo":             p.get("Activo", {}).get("checkbox", False),
        "perfil":             (p.get("Perfil", {}).get("rich_text") or [{}])[0].get("plain_text", ""),
        "rol":                (p.get("Rol objetivo", {}).get("rich_text") or [{}])[0].get("plain_text", ""),
        "stack":              [s["name"] for s in p.get("Stack", {}).get("multi_select", [])],
        "salario_min":        p.get("Salario min", {}).get("number", 0) or 0,
        "modalidad":          [m["name"] for m in p.get("Modalidad", {}).get("multi_select", [])],
        "ciudad":             (p.get("Ciudad", {}).get("rich_text") or [{}])[0].get("plain_text", ""),
        "telefono":           (p.get("Teléfono", {}).get("rich_text") or [{}])[0].get("plain_text", ""),
        "linkedin":           p.get("LinkedIn", {}).get("url", "") or "",
        "cv_master_url":      p.get("CV Master URL", {}).get("url", "") or "",
        "cv_master_url_es":   p.get("CV Master URL ES", {}).get("url", "") or "",
        "cv_master_file_id":  (p.get("cv_master_file_id", {}).get("rich_text") or [{}])[0].get("plain_text", ""),
    }


def buscar_oferta_en_notion(empresa: str, puesto: str) -> dict | None:
    """Busca una oferta en la DB de Ofertas por Empresa + Puesto.

    Sirve para recuperar datos que n8n guardó en Notion al crear la oferta
    pero que no llegan en el body de /generar-cv o /generar-carta (sobre todo
    la Descripción, clave para detectar el idioma). Devuelve dict o None.
    """
    if not NOTION_DB_OFERTAS or not empresa or not puesto:
        return None
    try:
        resp = requests.post(
            f"https://api.notion.com/v1/databases/{NOTION_DB_OFERTAS}/query",
            headers=notion_headers(),
            json={"filter": {"and": [
                {"property": "Empresa", "title": {"equals": empresa}},
                {"property": "Puesto", "rich_text": {"equals": puesto}},
            ]}, "page_size": 1},
            timeout=15,
        )
    except requests.RequestException as e:
        logger.warning("Notion query oferta falló: %s", e)
        return None
    if resp.status_code != 200:
        logger.warning("Notion query oferta error %s: %s", resp.status_code, resp.text[:200])
        return None
    results = resp.json().get("results", [])
    if not results:
        return None
    p = results[0].get("properties", {})
    return {
        "descripcion":     (p.get("Descripción", {}).get("rich_text") or [{}])[0].get("plain_text", ""),
        "nombre_contacto": (p.get("Nombre Contacto", {}).get("rich_text") or [{}])[0].get("plain_text", ""),
        "idioma":          (p.get("Idioma", {}).get("select") or {}).get("name", ""),
    }


def guardar_link_cv_en_notion(empresa: str, puesto: str, link_drive: str,
                              nombre_archivo: str) -> bool:
    """Escribe el enlace del CV en la ficha de la oferta, aqui y ahora.

    Nace el 30jul2026. Hasta hoy este enlace lo escribia n8n DESPUES de recibir
    la respuesta de /generar-cv, en una rama paralela a la del email. El nodo que
    llama aqui tiene timeout de 120s: con Sonnet la generacion se pasa de ahi,
    n8n aborta, y este servidor termina y sube el DOCX a Drive igualmente. El
    resultado es un CV huerfano — existe en Drive, y ni la ficha ni Veronica se
    enteran. Paso tres veces el mismo dia (Cactus, Alan, Trivelta).

    El enlace se escribe donde se sube el fichero para que la invariante sea
    cierta por construccion: si el CV existe, el enlace existe. Lo que haga n8n
    despues deja de importar para esto.

    Best-effort a proposito: si Notion falla, se loguea y se sigue. Un CV
    generado no se tira por no haber podido anotarlo.
    """
    if not NOTION_DB_OFERTAS or not empresa or not puesto or not link_drive:
        return False
    try:
        resp = requests.post(
            f"https://api.notion.com/v1/databases/{NOTION_DB_OFERTAS}/query",
            headers=notion_headers(),
            json={"filter": {"and": [
                {"property": "Empresa", "title": {"equals": empresa}},
                {"property": "Puesto", "rich_text": {"equals": puesto}},
            ]}, "page_size": 1},
            timeout=15,
        )
        if resp.status_code != 200:
            logger.warning("Notion: query oferta %s error %s", empresa, resp.status_code)
            return False
        results = resp.json().get("results", [])
        if not results:
            logger.warning("Notion: no se encontro la oferta %s / %s para anotar el CV",
                           empresa, puesto)
            return False

        page_id = results[0]["id"]
        patch = requests.patch(
            f"https://api.notion.com/v1/pages/{page_id}",
            headers=notion_headers(),
            json={"properties": {
                "Link CV Drive": {"url": link_drive},
                "CV usado": {"rich_text": [{"text": {"content": nombre_archivo[:2000]}}]},
            }},
            timeout=15,
        )
        if patch.status_code != 200:
            logger.warning("Notion: patch link CV error %s: %s",
                           patch.status_code, patch.text[:200])
            return False
        logger.info("Notion: enlace del CV anotado en %s / %s", empresa, puesto)
        return True
    except requests.RequestException as e:
        logger.warning("Notion: no se pudo anotar el enlace del CV: %s", e)
        return False


def _extraer_drive_file_id(url: str) -> str:
    """Extrae el file ID de una URL de Google Drive."""
    import re
    m = re.search(r'/d/([a-zA-Z0-9_-]+)', url) or re.search(r'id=([a-zA-Z0-9_-]+)', url)
    return m.group(1) if m else ""


def crear_usuario_en_notion(datos: dict) -> dict:
    """Crea un usuario en la BD de Notion."""
    url = "https://api.notion.com/v1/pages"
    cv_master_url = datos.get("cv_master_url") or ""
    cv_master_file_id = _extraer_drive_file_id(cv_master_url) if cv_master_url else ""
    props = {
        "Name":           {"title":  [{"text": {"content": datos.get("nombre", "")}}]},
        "Email":          {"email":   datos.get("email", "")},
        "Perfil":         {"rich_text": [{"text": {"content": datos.get("perfil", "")}}]},
        "Rol objetivo":   {"rich_text": [{"text": {"content": datos.get("rol_objetivo", "") or datos.get("rol", "")}}]},
        "Stack":          {"multi_select": [{"name": s} for s in datos.get("stack", [])]},
        "Salario min":    {"number": datos.get("salario_min") or datos.get("salario") or 0},
        "Modalidad":      {"multi_select": [{"name": m} for m in datos.get("modalidad", [])]},
        "Ciudad":         {"rich_text": [{"text": {"content": datos.get("ciudad", "")}}]},
        "LinkedIn":       {"url": datos.get("linkedin") or None},
        "CV Master URL":  {"url": cv_master_url or None},
        "Activo":         {"checkbox": True},
    }
    if cv_master_file_id:
        props["cv_master_file_id"] = {"rich_text": [{"text": {"content": cv_master_file_id}}]}
    # Filtrar propiedades con valor None que Notion rechaza
    payload = {
        "parent": {"database_id": NOTION_DB_USUARIOS},
        "properties": {k: v for k, v in props.items() if v is not None and v != {"url": None}},
    }
    resp = requests.post(url, headers=notion_headers(), json=payload, timeout=15)
    resp.raise_for_status()
    return resp.json()


def crear_oferta_en_notion(oferta: dict, idioma: str = "", usuario_notion_id: str = "") -> dict:
    """Crea una oferta en la DB Ofertas con TODOS sus campos mapeados.

    Centraliza el mapeo oferta→Notion en código (en vez de la UI de n8n).
    `oferta` es el dict que devuelve /buscar-ofertas-reales (real_jobs.py).
    """
    def _txt(v: str) -> dict:
        return {"rich_text": [{"text": {"content": (v or "")[:2000]}}]}

    props = {
        "Empresa":       {"title": [{"text": {"content": oferta.get("empresa", "")}}]},
        "Puesto":        _txt(oferta.get("puesto", "")),
        "Descripción":   _txt(oferta.get("descripcion", "")),
        "Salario":       _txt(oferta.get("salario", "")),
        "Ubicación":     _txt(oferta.get("ubicacion", "")),
        "Tipo Contrato": _txt(oferta.get("tipo_contrato", "")),
        "Estado":        {"select": {"name": "Pendiente"}},
    }
    modalidad = oferta.get("modalidad", "")
    if modalidad in ("Remoto", "Hibrido", "Presencial"):
        props["Modalidad"] = {"select": {"name": modalidad}}
    tags = [t for t in (oferta.get("tags") or []) if t]
    if tags:
        props["Tags"] = {"multi_select": [{"name": t[:100]} for t in tags]}
    if oferta.get("link"):
        props["Link oferta"] = {"url": oferta["link"]}
    if oferta.get("fecha_publicacion"):
        props["Fecha Publicacion"] = {"date": {"start": oferta["fecha_publicacion"]}}
    if idioma in ("en", "es"):
        props["Idioma"] = {"select": {"name": idioma}}
    if usuario_notion_id:
        props["Usuario"] = {"relation": [{"id": usuario_notion_id}]}

    payload = {"parent": {"database_id": NOTION_DB_OFERTAS}, "properties": props}
    resp = requests.post("https://api.notion.com/v1/pages",
                         headers=notion_headers(), json=payload, timeout=15)
    resp.raise_for_status()
    return resp.json()


# ══════════════════════════════════════════════
# GENERACIÓN DOCX
# ══════════════════════════════════════════════

# Guiones largos/medios y flechas: rastro tipográfico de IA. Regla NO NEGOCIABLE
# de la usuaria — jamás deben aparecer en un CV o carta que sale a una empresa.
_ARROWS = "→←⟶⟹➜➔➡⇒"
_DASHES = "—–―‒−"
_RE_ARROW = re.compile(r"\s*[" + _ARROWS + r"]\s*")
_RE_DASH  = re.compile(r"\s*[" + _DASHES + r"]\s*")
_RE_SPACES = re.compile(r"[ \t]{2,}")


def sanear_tipografia(texto: str, idioma: str = "es") -> str:
    """Elimina guiones largos/medios (—, –) y flechas (→) del texto final.

    Se aplica en el RENDER (DOCX y carta), nunca sobre el texto que el parser del
    DOCX usa para detectar estructura (ese sigue viendo el — crudo). Las flechas se
    traducen a la palabra de transición del idioma ("a"/"to"); los guiones a guion
    normal. Es una red determinista: no depende de que el LLM obedezca el prompt."""
    if not texto:
        return texto
    trans = " to " if idioma == "en" else " a "
    t = _RE_ARROW.sub(trans, texto)
    t = _RE_DASH.sub(" - ", t)
    t = _RE_SPACES.sub(" ", t)
    return t


def generar_docx(contenido_cv: str, nombre_candidato: str) -> bytes:
    """Wrapper legacy — usar generar_docx_con_cabecera() para nuevos CVs."""
    return generar_docx_con_cabecera(contenido_cv, {"nombre": nombre_candidato})


def generar_docx_con_cabecera(contenido_cv: str, usuario: dict, titular: str = "", idioma: str = "es") -> bytes:
    """Genera DOCX con cabecera profesional estructurada usando datos reales del usuario.
    `titular` (si viene) es el headline adaptado a la oferta por el LLM; tiene prioridad
    sobre el campo `rol` fijo del perfil."""
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLUE = RGBColor(0x1A, 0x56, 0xDB)
    DARK = RGBColor(0x1A, 0x1A, 0x1A)
    GREY = RGBColor(0x66, 0x66, 0x66)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # ── Cabecera ──────────────────────────────────────────────────
    nombre   = usuario.get("nombre", "Candidato")
    rol      = titular or usuario.get("rol", "")
    ciudad   = usuario.get("ciudad", "")
    telefono = usuario.get("telefono", "")
    # Email de cabecera (contacto) separado del email-clave de búsqueda en Notion
    email    = usuario.get("email_cv") or usuario.get("email", "")
    linkedin = (usuario.get("linkedin", "") or "").replace("https://", "").replace("http://", "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(nombre.upper())
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = DARK

    if rol:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(sanear_tipografia(rol, idioma))
        r.font.size = Pt(11); r.font.color.rgb = BLUE

    contacto = " · ".join(c for c in [ciudad, telefono, email, linkedin] if c)
    if contacto:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(contacto)
        r.font.size = Pt(8.5); r.font.color.rgb = GREY

    # Línea separadora
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1A56DB")
    pBdr.append(bottom); pPr.append(pBdr)

    # ── Cuerpo del CV ────────────────────────────────────────────
    SECCIONES = ["PERFIL PROFESIONAL", "EXPERIENCIA PROFESIONAL", "EXPERIENCIA",
                 "HABILIDADES TÉCNICAS", "HABILIDADES", "FORMACIÓN", "IDIOMAS",
                 "PROYECTOS", "CERTIFICACIONES", "COMPETENCIAS"]

    for linea in contenido_cv.strip().split("\n"):
        linea = linea.strip()
        if not linea:
            continue

        limpia = linea.upper().strip()
        # La DETECCIÓN de estructura usa `linea` cruda (necesita ver el — como
        # marcador de empresa). El RENDER usa la versión saneada: así ningún
        # guion largo ni flecha llega nunca al DOCX. Regla NO NEGOCIABLE.
        render = sanear_tipografia(linea, idioma)

        # Sección
        if any(limpia.startswith(s) for s in SECCIONES) and len(linea) < 50:
            p = doc.add_paragraph()
            r = p.add_run(render.upper())
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLUE
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1A56DB")
            pBdr.append(bottom); pPr.append(pBdr)
            continue

        # Bullet
        if linea.startswith(("- ", "• ", "* ")):
            p = doc.add_paragraph()
            r = p.add_run("• " + render[2:].strip())
            r.font.size = Pt(9.5); r.font.color.rgb = DARK
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after  = Pt(2)
            continue

        # Empresa / puesto (línea con — o –)
        if ("—" in linea or "–" in linea) and len(linea) < 100:
            p = doc.add_paragraph()
            r = p.add_run(render)
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(1)
            continue

        # Fecha (línea corta con año)
        if re.search(r"(20\d{2}|19\d{2})", linea) and len(linea) < 60:
            p = doc.add_paragraph()
            r = p.add_run(render)
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
            p.paragraph_format.space_after = Pt(2)
            continue

        # Texto normal
        p = doc.add_paragraph()
        r = p.add_run(render)
        r.font.size = Pt(9.5); r.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(3)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════
# FORMULARIO HTML — 3 pantallas (email → existente | nuevo → completo)
# ══════════════════════════════════════════════

FORMULARIO_HTML = """
<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="UTF-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
  <title>BuscarTrabajo — Registro</title>
  <style>
    * { box-sizing: border-box; margin: 0; padding: 0; }
    body { font-family: 'Segoe UI', sans-serif; background: #f0f4ff; display: flex;
           justify-content: center; align-items: flex-start; min-height: 100vh; padding: 2rem 1rem; }
    .card { background: #fff; border-radius: 16px; padding: 2rem; max-width: 520px;
            width: 100%; box-shadow: 0 4px 24px rgba(0,0,0,.08); }
    h1 { font-size: 1.5rem; color: #1a56db; margin-bottom: .25rem; }
    .sub { color: #6b7280; font-size: .9rem; margin-bottom: 1.5rem; }
    label { display: block; font-size: .85rem; color: #374151; margin-bottom: .25rem; font-weight: 500; }
    input, textarea, select { width: 100%; padding: .6rem .8rem; border: 1px solid #d1d5db;
      border-radius: 8px; font-size: .95rem; margin-bottom: 1rem; }
    textarea { resize: vertical; min-height: 80px; }
    .screen { display: none; }
    .screen.active { display: block; }
    button { width: 100%; padding: .75rem; background: #1a56db; color: #fff;
             border: none; border-radius: 8px; font-size: 1rem; cursor: pointer; font-weight: 600; }
    button:hover { background: #1648c0; }
    button:disabled { background: #9ca3af; cursor: not-allowed; }
    button.secondary { background: #22c55e; }
    button.secondary:hover { background: #16a34a; }
    button.outline { background: #fff; color: #1a56db; border: 2px solid #1a56db; }
    button.outline:hover { background: #f0f7ff; }
    .button-row { display: flex; gap: .75rem; }
    .msg { margin-top: 1rem; padding: .75rem; border-radius: 8px; font-size: .9rem; }
    .ok  { background: #d1fae5; color: #065f46; }
    .err { background: #fee2e2; color: #991b1b; }
    .step { color: #9ca3af; font-size: .8rem; margin-bottom: 1rem; }
    .link-usuarios { text-align: center; margin-bottom: 1rem; }
    .link-usuarios a { color: #1a56db; font-size: .85rem; text-decoration: none; }
  </style>
</head>
<body>
<div class="card">
  <div class="link-usuarios">
    <a href="/usuarios" target="_blank">📋 Ver usuarios registrados</a>
  </div>

  <!-- PANTALLA 1a — Solo email (check si existe) -->
  <div id="sEmail" class="screen active">
    <h1>🚀 BuscarTrabajo.ai</h1>
    <p class="sub">Te buscamos trabajo mientras duermes.</p>
    <label>Email</label>
    <input id="emailInicial" type="email" placeholder="tu@email.com" />
    <button type="button" onclick="comprobarEmail()">Continuar →</button>
    <div id="msgEmail"></div>
  </div>

  <!-- PANTALLA 2a — Usuario existente -->
  <div id="sExistente" class="screen">
    <h1 id="saludoExistente">¡Hola de nuevo!</h1>
    <p class="sub">¿Cuándo quieres que busquemos ofertas?</p>
    <div class="button-row">
      <button class="secondary" onclick="accionExistente('ahora')">⚡ Buscar ahora</button>
      <button class="outline" onclick="accionExistente('manana')">🌅 Mañana a las 9</button>
    </div>
    <div id="msgExistente"></div>
  </div>

  <!-- PANTALLA 1 — Datos básicos (usuario nuevo) -->
  <div id="s1" class="screen">
    <h1>🎯 Cuéntanos qué buscas</h1>
    <p class="sub">Solo una vez — luego te buscamos ofertas cada día.</p>
    <p class="step">Paso 1 de 2</p>
    <label>Nombre completo</label>
    <input id="nombre" placeholder="Ana García López" />
    <label>Email</label>
    <input id="email" type="email" readonly style="background:#f0f0f0;color:#666;" />
    <label>Perfil profesional <span style="color:#9ca3af">(breve descripción)</span></label>
    <textarea id="perfil" placeholder="Desarrolladora frontend con 5 años de experiencia en React y Vue…"></textarea>
    <button type="button" onclick="irS2()">Continuar →</button>
  </div>

  <!-- PANTALLA 2 — Preferencias + Buscar ahora -->
  <div id="s2" class="screen">
    <p class="step">Paso 2 de 2</p>
    <label>Rol objetivo</label>
    <input id="rol" placeholder="Senior Frontend Developer" />
    <label>Stack principal <span style="color:#9ca3af">(separado por comas)</span></label>
    <input id="stack" placeholder="React, TypeScript, Node.js" />
    <label>Salario mínimo (€ bruto/año)</label>
    <input id="salario" type="number" placeholder="40000" />
    <label>Modalidad</label>
    <select id="modalidad">
      <option value="Remoto">Remoto</option>
      <option value="Híbrido">Híbrido</option>
      <option value="Presencial">Presencial</option>
    </select>
    <label>Ciudad (si aplica)</label>
    <input id="ciudad" placeholder="Madrid, Barcelona…" />
    <label>LinkedIn <span style="color:#9ca3af">(opcional)</span></label>
    <input id="linkedin" placeholder="https://linkedin.com/in/tu-perfil" />
    <label>CV Master (link Google Drive, opcional)</label>
    <input id="cv_master_url" placeholder="https://drive.google.com/file/d/..." />
    <button type="button" onclick="registrar()">🔍 Registrarme y buscar ahora</button>
    <div id="msg"></div>
  </div>

  <!-- PANTALLA 3 — Listo -->
  <div id="sListo" class="screen">
    <h1>✅ ¡Listo!</h1>
    <p class="sub" id="confirmacion">Todo en orden.</p>
  </div>
</div>

<script>
let currentEmail = '';
let currentNombre = '';

function showScreen(id) {
  document.querySelectorAll('.screen').forEach(s => s.classList.remove('active'));
  document.getElementById(id).classList.add('active');
}

// PANTALLA 1a — comprobar email
async function comprobarEmail() {
  const email = document.getElementById('emailInicial').value.trim();
  const msg = document.getElementById('msgEmail');
  if (!email) {
    msg.innerHTML = '<div class="msg err">Introduce un email válido</div>';
    return;
  }
  currentEmail = email;
  const btn = document.querySelector('#sEmail button');
  btn.disabled = true;
  btn.textContent = 'Comprobando…';

  try {
    const resp = await fetch('/check-email', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ email })
    });
    const data = await resp.json();

    if (data.existe) {
      currentNombre = data.nombre || '';
      document.getElementById('saludoExistente').textContent = `¡Hola de nuevo, ${data.nombre || ''}!`;
      showScreen('sExistente');
    } else {
      document.getElementById('email').value = email;
      showScreen('s1');
    }
  } catch(e) {
    msg.innerHTML = '<div class="msg err">Error: ' + e.message + '</div>';
    btn.disabled = false;
    btn.textContent = 'Continuar →';
  }
}

// PANTALLA 2a — usuario existente
async function accionExistente(accion) {
  const msg = document.getElementById('msgExistente');
  try {
    const resp = await fetch('/accion-existente', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ email: currentEmail, nombre: currentNombre, accion })
    });
    const data = await resp.json();
    if (accion === 'ahora') {
      // La respuesta se MIRA. Antes se tiraba y esta pantalla cantaba exito
      // aunque n8n no hubiera recibido nada.
      document.getElementById('confirmacion').textContent = data.busqueda_disparada
        ? 'Buscando ahora mismo. Recibirás las ofertas en unos minutos en tu email.'
        : 'No se ha podido lanzar la búsqueda ahora mismo. Entras igualmente en el barrido de las 9:00.';
    } else {
      document.getElementById('confirmacion').textContent =
        'De acuerdo. Mañana a las 9:00 recibirás tus ofertas personalizadas.';
    }
    showScreen('sListo');
  } catch(e) {
    msg.innerHTML = '<div class="msg err">Error: ' + e.message + '</div>';
  }
}

// PANTALLA 1 → 2 (usuario nuevo)
function irS2() {
  if (!document.getElementById('nombre').value.trim()) {
    alert('Por favor rellena el nombre.');
    return;
  }
  document.getElementById('s1').classList.remove('active');
  document.getElementById('s2').classList.add('active');
}

// PANTALLA 2 — registrar nuevo
async function registrar() {
  const btn = document.querySelector('#s2 button');
  btn.disabled = true;
  btn.textContent = 'Procesando…';
  const msg = document.getElementById('msg');
  msg.innerHTML = '';

  const payload = {
    nombre:        document.getElementById('nombre').value.trim(),
    email:         document.getElementById('email').value.trim(),
    perfil:        document.getElementById('perfil').value.trim(),
    rol_objetivo:  document.getElementById('rol').value.trim(),
    stack:         document.getElementById('stack').value.split(',').map(s=>s.trim()).filter(Boolean),
    salario_min:   parseInt(document.getElementById('salario').value) || 0,
    modalidad:     [document.getElementById('modalidad').value],
    ciudad:        document.getElementById('ciudad').value.trim(),
    linkedin:      document.getElementById('linkedin').value.trim(),
    cv_master_url: document.getElementById('cv_master_url').value.trim(),
  };

  try {
    const resp = await fetch('/registro', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify(payload),
    });
    const data = await resp.json();
    if (data.ok) {
      document.getElementById('confirmacion').textContent =
        data.mensaje || '¡Registro completado! En breve recibirás ofertas.';
      showScreen('sListo');
    } else {
      msg.innerHTML = '<div class="msg err">❌ ' + (data.error || 'Error inesperado') + '</div>';
      btn.disabled = false; btn.textContent = '🔍 Registrarme y buscar ahora';
    }
  } catch(e) {
    msg.innerHTML = '<div class="msg err">❌ Error de conexión: ' + e.message + '</div>';
      btn.disabled = false; btn.textContent = '🔍 Registrarme y buscar ahora';
  }
}
</script>
</body>
</html>
"""


# ══════════════════════════════════════════════
# ENDPOINTS
# ══════════════════════════════════════════════

@app.route("/")
def index():
    return render_template_string(FORMULARIO_HTML)


@app.route("/health")
def health():
    return jsonify({
        "status":       "ok",
        "version":      "v2.4",
        # QUE MODELO escribe cada cosa, leido de las variables reales. Antes esto
        # devolvia las constantes "groq" y "v2.3-groq", que era falso: Groq es solo el
        # fallback si Claude falla. Sin esto no se puede verificar desde fuera un
        # cambio de CV_MODEL en Render, que es justo cuando mas falta hace.
        "modelos":      {
            "cv":       CV_MODEL,
            "carta":    CARTA_MODEL,
            "fallback": GROQ_MODEL,
        },
        "fallbacks":    {
            "gemini":  bool(GEMINI_API_KEY),
            "claude":  bool(CLAUDE_API_KEY),
        },
        "deploy":       {
            "branch": os.environ.get("RENDER_GIT_BRANCH", "local"),
            "commit": (os.environ.get("RENDER_GIT_COMMIT", "") or "")[:7],
        },
        "timestamp":    datetime.now(timezone.utc).isoformat(),
    })


@app.route("/debug")
def debug():
    """Prueba rápida del LLM activo (Groq primero)."""
    try:
        r = call_llm("Responde solo: 'Groq funcionando correctamente en cv_server v2.3'")
        return jsonify({"ok": True, "respuesta": r.contenido, "modelo": r.modelo})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/check-email", methods=["POST"])
def check_email():
    """Comprueba si un email ya existe en Notion. Devuelve {existe, nombre}."""
    datos = request.get_json(force=True)
    email = (datos.get("email") or "").strip().lower()
    if not email:
        return jsonify({"existe": False, "error": "email requerido"}), 400

    try:
        usuario = buscar_usuario_por_email(email)
    except Exception as e:
        logger.error("Error check-email: %s", e)
        return jsonify({"existe": False, "error": str(e)}), 500

    if usuario and usuario.get("activo"):
        return jsonify({
            "existe": True,
            "nombre": usuario.get("nombre", ""),
            "email":  email,
        })
    return jsonify({"existe": False, "email": email})


def payload_buscar_para_user(usuario: dict) -> dict:
    """Traduce un usuario al contrato que espera el webhook `buscar-para-user`.

    El contrato no es invento: es el que produce el nodo `Code — Normalizar users
    (schedule)` del workflow de PROD, que es por donde entra el disparo de las 9:00.
    Dos nombres NO coinciden con los del perfil que devuelve Notion, y por eso
    existe esta funcion: `Salario min` viaja como `salario`, y el id de la pagina
    como `user_id`.

    Acepta las dos formas del usuario: la que devuelve `buscar_usuario_por_email`
    y la cruda del formulario de alta (`rol_objetivo`, `salario`).
    """
    u = usuario or {}
    return {
        "user_id":       u.get("notion_id") or u.get("user_id") or "",
        "nombre":        u.get("nombre", ""),
        "email":         u.get("email", ""),
        "email_usuario": u.get("email", ""),
        "perfil":        u.get("perfil", ""),
        "rol":           u.get("rol") or u.get("rol_objetivo") or "",
        "stack":         u.get("stack") or [],
        "salario":       u.get("salario_min") or u.get("salario") or 0,
        "modalidad":     u.get("modalidad") or [],
        "ciudad":        u.get("ciudad", ""),
        "linkedin":      u.get("linkedin", ""),
        "cv_master_url": u.get("cv_master_url", ""),
        "source":        "cv-server",
    }


def disparar_busqueda(usuario: dict) -> bool:
    """Pide a n8n una búsqueda para este usuario. Devuelve si n8n la aceptó.

    Devuelve un bool a propósito. Antes esto era fire & forget con un `warning`
    y quien llamaba no tenía forma de saber si había pasado algo, así que la API
    respondía `ok: true` con el webhook devolviendo 404. Un 404 aquí significa
    que el path no está dado de alta en n8n: es un fallo de configuración que hay
    que ver, no un aviso que se traga el log.
    """
    if not usuario or not WEBHOOK_BUSCAR_AHORA:
        logger.error("Búsqueda NO disparada: falta usuario o WEBHOOK_BUSCAR_AHORA")
        return False
    try:
        r = requests.post(WEBHOOK_BUSCAR_AHORA, json=payload_buscar_para_user(usuario), timeout=8)
    except Exception as e:
        logger.error("Búsqueda NO disparada, n8n no responde (%s): %s", WEBHOOK_BUSCAR_AHORA, e)
        return False
    if r.status_code >= 400:
        logger.error(
            "Búsqueda NO disparada, n8n devolvió %s en %s. Un 404 aquí = el webhook "
            "no existe en la instancia.", r.status_code, WEBHOOK_BUSCAR_AHORA,
        )
        return False
    logger.info("Búsqueda disparada para %s", usuario.get("email", ""))
    return True


@app.route("/accion-existente", methods=["POST"])
def accion_existente():
    """Usuario existente pulsa 'Buscar ahora' o 'Mañana 9am'."""
    datos = request.get_json(force=True)
    email = (datos.get("email") or "").strip().lower()
    accion = datos.get("accion", "")

    if not email:
        return jsonify({"ok": False, "error": "email requerido"}), 400

    disparada = False
    if accion == "ahora":
        # El perfil hace falta ENTERO: mandando solo email y nombre, n8n buscaba
        # ofertas sin rol, sin stack y sin salario, o sea para nadie.
        try:
            usuario = buscar_usuario_por_email(email)
        except Exception as e:
            logger.error("No se pudo leer el usuario %s en Notion: %s", email, e)
            usuario = None
        disparada = disparar_busqueda(usuario)

    return jsonify({
        "ok": True,
        "accion": accion,
        "email": email,
        "busqueda_disparada": disparada,
    })


@app.route("/registro", methods=["POST"])
def registro():
    """Registra usuario nuevo en Notion y dispara webhook n8n."""
    datos = request.get_json(force=True)
    email = (datos.get("email") or "").strip().lower()
    if not email:
        return jsonify({"ok": False, "error": "email requerido"}), 400

    # Si ya existe, no duplicar
    try:
        existente = buscar_usuario_por_email(email)
    except Exception:
        existente = None

    if existente:
        # Usuario ya existe → disparar búsqueda igual
        disparada = disparar_busqueda(existente)
        return jsonify({
            "ok": True,
            "mensaje": ("Ya estabas registrado. Buscando ofertas ahora mismo."
                        if disparada else
                        "Ya estabas registrado. No se pudo lanzar la búsqueda ahora; "
                        "entrarás en el barrido de las 9:00."),
            "email": email,
            "busqueda_disparada": disparada,
        })

    # Crear en Notion
    try:
        notion_page = crear_usuario_en_notion(datos)
        notion_id = notion_page.get("id", "")
    except Exception as e:
        logger.error("Notion error: %s", e)
        return jsonify({"ok": False, "error": f"Error creando usuario en Notion: {e}"}), 500

    # El alta ya está hecha en Notion. Esto lanza la primera búsqueda: iba al
    # webhook `nuevo-usuario`, que no existe, así que quien se registraba no
    # recibía nada hasta el barrido de las 9:00 del día siguiente.
    disparada = disparar_busqueda({**datos, "notion_id": notion_id})

    return jsonify({
        "ok":      True,
        "mensaje": ("Usuario registrado. Buscando tus primeras ofertas ahora mismo."
                    if disparada else
                    "Usuario registrado. Recibirás ofertas en el barrido de las 9:00."),
        "email":   email,
        "busqueda_disparada": disparada,
    })


def generar_cv_core(email: str, empresa: str, puesto: str,
                    descripcion: str = "", idioma_in: str = "") -> dict:
    """Núcleo de /generar-cv (sin Flask): orquesta Notion/Drive/LLM y devuelve
    el dict de respuesta. Lanza CVError(status, msg) en los errores. Lo comparten
    la ruta Flask y la ruta FastAPI. Ver docs/ADR-001."""
    if not email or not empresa or not puesto:
        raise CVError(400, "email, empresa y puesto son requeridos")

    # Idioma autoritativo: el que n8n decidió de la descripción REAL al crear la
    # oferta (la de Notion suele ser un resumen en español que confunde la
    # detección). Prioridad: body.idioma > Idioma guardado en Notion > detección.
    idioma_in = (idioma_in or "").strip().lower()
    if not descripcion.strip() or idioma_in not in ("en", "es"):
        oferta = buscar_oferta_en_notion(empresa, puesto)
        if oferta:
            descripcion = descripcion or oferta.get("descripcion", "")
            idioma_in = idioma_in or (oferta.get("idioma") or "").strip().lower()

    # 1. Leer perfil completo del usuario desde Notion
    usuario = buscar_usuario_por_email(email)
    if not usuario:
        raise CVError(404, f"Usuario {email} no encontrado en Notion")

    nombre = usuario.get("nombre") or email.split("@")[0]

    # 2. Resolver idioma y leer el CV master en ese idioma
    idioma = idioma_in if idioma_in in ("en", "es") else idioma_de_oferta(puesto, descripcion, empresa)
    tiene_master_configurado = _tiene_algun_master(usuario)
    try:
        master     = leer_cv_master_desde_drive(usuario, idioma)
        cv_master  = master.texto
    except Exception as e:
        # get_drive_service() -> creds.refresh() puede fallar (token caducado)
        # ANTES del try interno de la lectura: sin este guard, sale un 500 HTML
        # opaco. Devolvemos el error REAL en JSON para poder diagnosticar.
        logger.error("Drive auth/lectura falló en /generar-cv: %s", e)
        raise CVError(502, f"No se pudo autenticar/leer el CV master en Drive: {e}")

    # Guardrail: si hay un master configurado pero lo leído es ilegible
    # (basura binaria de un .docx sin parsear, o sin acceso), NO generamos
    # nada. Mejor fallar claro que mandar un CV con datos inventados.
    def _es_legible(t: str) -> bool:
        if not t:
            return False
        if t.lstrip().startswith("PK"):  # firma ZIP de un .docx no parseado
            return False
        imprimibles = sum(1 for c in t if c.isprintable() or c in "\n\r\t")
        return imprimibles / len(t) >= 0.85

    if tiene_master_configurado and not _es_legible(cv_master):
        logger.error("CV master ILEGIBLE para %s (largo=%d) — abortando para no inventar",
                     email, len(cv_master or ""))
        raise CVError(502, "No se pudo leer tu CV master desde Drive (archivo ilegible o sin acceso). "
                           "NO se generó un CV para evitar enviar datos inventados. "
                           "Revisá el archivo y los permisos en Drive.")

    if cv_master:
        logger.info("CV master leído (%d chars) para %s", len(cv_master), email)
    else:
        logger.warning("CV master no encontrado para %s — usando solo perfil de Notion", email)

    # 3. Construir contexto del candidato
    ciudad = usuario.get("ciudad", "Madrid")
    rol    = usuario.get("rol", "")
    stack  = ", ".join(usuario.get("stack", [])) or "React, TypeScript"

    # 4. Prompt con 4 fases + CV master real
    if cv_master:
        contexto_candidato = f"""CV MASTER COMPLETO (usa SOLO esta experiencia, NO inventes):
{cv_master}"""
    else:
        contexto_candidato = f"""PERFIL DEL CANDIDATO (sin CV master disponible):
- Nombre: {nombre}
- Rol objetivo: {rol}
- Stack: {stack}
- Ciudad: {ciudad}
- Perfil: {usuario.get("perfil", "")}"""

    # Titulos de seccion y regla de idioma, en el idioma de la oferta
    idioma_nombre = "English" if idioma == "en" else "Spanish"
    if idioma == "en":
        bloque_formato = """OUTPUT FORMAT (plain text, no markdown):

HEADLINE: [professional title for this offer — see HEADLINE RULES below]

PROFESSIONAL SUMMARY
[2 full paragraphs (4-6 lines each) tailored to the offer, generated from her REAL EXPERIENCE (never copied from PERFIL BASE). Write with NO GRAMMATICAL SUBJECT, standard English CV style: "Frontend Tech Lead with 10+ years...", "Led the migration...". NEVER "She is", "She brings", "Her career spans", and never "I am". First paragraph: the profile itself + core strengths relevant to this role. Second paragraph: depth, domains and the angle that fits this offer.]

PROFESSIONAL EXPERIENCE
[Role] — [Company]
[Start date] - [End date]
- Real achievement from the CV master, XYZ formula, prioritised by relevance
- Real achievement from the CV master, XYZ formula, prioritised by relevance
- Real achievement from the CV master, XYZ formula, prioritised by relevance
- Real achievement from the CV master, XYZ formula, prioritised by relevance
- Real achievement from the CV master, XYZ formula, prioritised by relevance
- Real achievement from the CV master, XYZ formula, prioritised by relevance
(6-9 bullets for recent/relevant roles, 3-4 for older ones — always real, never padded)

TECHNICAL SKILLS
[Skills grouped by category (Frontend, AI, Design Systems, Backend, Cloud, Testing...) with concrete tools/versions, ordered by relevance to this offer]

EDUCATION
[From the CV master]

LANGUAGES
[From the CV master]

FINAL RULES:
- First line MUST be "HEADLINE: ..." — it becomes the header title
- Do NOT include name/email/phone, they are added programmatically
- Do NOT use markdown (**text**, ##, ```)
- Do NOT invent anything not in the CV master
- EXPERIENCE reads as a career story told through ROLES: the job title opens every entry and the company follows it. A recruiter must be able to scan the left edge and see the progression (Tech Lead, then Front-End Developer, then Designer). Never lead with the company.
- Language: the ENTIRE CV must be in English (section titles and content)"""
    else:
        bloque_formato = """FORMATO DE SALIDA (texto plano, sin markdown):

HEADLINE: [titular profesional para esta oferta — ver REGLAS DEL HEADLINE abajo]

PERFIL PROFESIONAL
[2 párrafos completos (4-6 líneas cada uno) adaptados a la oferta, generados desde su EXPERIENCIA real (NUNCA copiados del PERFIL BASE). Primer párrafo: quién es + fortalezas clave relevantes para este puesto. Segundo párrafo: profundidad, dominios y el ángulo que encaja con esta oferta.]

EXPERIENCIA PROFESIONAL
[Puesto] — [Empresa]
[Fecha inicio] - [Fecha fin]
- Logro real del CV master, fórmula XYZ, priorizado por relevancia
- Logro real del CV master, fórmula XYZ, priorizado por relevancia
- Logro real del CV master, fórmula XYZ, priorizado por relevancia
- Logro real del CV master, fórmula XYZ, priorizado por relevancia
- Logro real del CV master, fórmula XYZ, priorizado por relevancia
- Logro real del CV master, fórmula XYZ, priorizado por relevancia
(6-9 bullets en los puestos recientes/relevantes, 3-4 en los antiguos — siempre reales, nunca de relleno)

HABILIDADES TÉCNICAS
[Skills agrupadas por categoría (Frontend, IA, Sistemas de Diseño, Backend, Cloud, Testing...) con herramientas/versiones concretas, ordenadas por relevancia para esta oferta]

FORMACIÓN
[Del CV master]

IDIOMAS
[Del CV master]

REGLAS FINALES:
- La primera línea DEBE ser "HEADLINE: ..." — se usa como titular de la cabecera
- NO incluir nombre/email/tel, se añaden programáticamente
- NO usar markdown (**texto**, ##, ```)
- NO inventar nada que no esté en el CV master
- La EXPERIENCIA se lee como una trayectoria contada a través de los PUESTOS: el puesto abre cada entrada y la empresa va detrás. Quien lee debe poder recorrer el margen izquierdo y ver la progresión (Tech Lead, antes Front-End Developer, antes Diseñadora). Nunca empieces por la empresa.
- Idioma: TODO el CV en español (títulos de sección y contenido)"""

    prompt = f"""Act as a senior tech recruiter who screens 200+ CVs daily. Adapt this candidate's CV for a specific job offer.

The target job offer is written in {idioma_nombre}. Generate the ENTIRE CV in {idioma_nombre} — both the section titles and the content.

PRINCIPIO FUNDAMENTAL: el CV generado NO añade información nueva. Únicamente reorganiza, prioriza y redacta de forma distinta información YA existente y demostrable en el CV Master. Nunca al revés.

REGLA MAESTRA — PROYECCIÓN, NO IDENTIDAD NUEVA (de esta se derivan casi todas las demás): la adaptación debe producir una PROYECCIÓN distinta de la MISMA trayectoria profesional, NUNCA una nueva identidad profesional. El CV de esta oferta y el que se generó para cualquier otra tienen que leerse como la misma persona enfocando distinto, no como profesionales diferentes. Un recruiter que viera tres CV suyos debe reconocer a la misma candidata adaptando el contenido. Si un cambio la hace parecer otra profesional, ese cambio está MAL aunque cada frase por separado sea cierta. Consecuencias directas: no cambies el titular de forma radical, no subas el seniority, no inventes herramientas, no muevas skills a experiencia, no conviertas un proyecto propio en una organización grande. Solo cambia el ÉNFASIS.

JERARQUÍA DE FUENTES (respétala siempre):
- La OFERTA decide QUÉ enfatizar.
- El PERFIL BASE decide DESDE QUÉ IDENTIDAD se responde (quién es la candidata). Es una GUÍA de identidad y coherencia, NO una fuente de contenido, y NUNCA constituye evidencia. No copies literalmente sus frases al CV.
- La EXPERIENCIA (junto con proyectos, formación y conocimientos técnicos) es la ÚNICA fuente de evidencia: decide QUÉ se puede afirmar.
- El CV generado es solo una reorganización de esa evidencia según la oferta.
Ninguna afirmación entra en el CV si no está respaldada por la EXPERIENCIA, proyectos, formación o skills del Master; el PERFIL BASE por sí solo NO basta como evidencia.

{contexto_candidato}

OFERTA TARGET:
- Empresa: {empresa}
- Puesto: {puesto}
- Descripción: {descripcion or "No disponible"}

PASO 1 — ANÁLISIS INTERNO (SOLO mental — NO lo escribas en la respuesta):
Piensa, SIN volcarlo al output, en:
- Skills del CV master que encajan con esta oferta
- Keywords de la oferta que deben aparecer
- Logros que mejor demuestran el fit
- NO inventar experiencia, métricas ni logros
Tu respuesta DEBE empezar EXACTAMENTE con la línea "HEADLINE: ...". Prohibido escribir
análisis, títulos, encabezados o cualquier texto ANTES de esa línea.

PASO 2 — CV ADAPTADO (output principal):
Genera el CV adaptado con estas reglas ESTRICTAS:
1. NO INVENTAR NUNCA: solo experiencia real del CV master. Nada de tecnologías no usadas, responsabilidades no ejercidas, liderazgo de personas o arquitectura que no haya hecho, ni métricas/impacto exagerados. REGLA DE EVIDENCIA: una tecnología o skill SOLO puede aparecer si está respaldada en el Master por experiencia profesional, un proyecto o formación significativa; no basta con haberla tocado puntualmente.
2. Adapta el ORDEN y ÉNFASIS según la oferta, no el contenido
3. Optimización ATS: integra las palabras clave EXACTAS de la oferta cuando formen parte de su experiencia real; el CV debe quedar 100% defendible en entrevista
4. Bullets con fórmula XYZ ("Logré X, medido por Y, haciendo Z") SOLO cuando la cifra de "Y" esté LITERALMENTE en el CV Master. Si el Master no da una cifra, escribe el bullet SIN métrica (X + Z): describe qué construiste y cómo, sin cuantificar. Nada de bullets genéricos tipo "responsable de...".
4bis. PROHIBICIÓN DE CIFRAS (no negociable): no introduzcas NINGÚN número que no aparezca en el CV Master. Ni usuarios, ni porcentajes, ni ingresos, ni tamaños de equipo, ni volúmenes. Prohibido también cuantificar con palabras ("millones de", "miles de", "cientos de", "millions of", "thousands of") si esa magnitud no está en el Master. Ante la duda, describe sin número: un CV sin cifras es defendible en entrevista; uno con una cifra inventada, no.
5. Densidad real: NO recortes ni resumas el CV master. Los puestos recientes/relevantes deben llevar 6-9 bullets; los antiguos 3-4. Si el master tiene el detalle, úsalo entero.
6. Redacta como PERFIL DE PRODUCTO: traducción de necesidades de negocio a soluciones digitales, colaboración con diseño y producto, plataformas B2B y B2C, diseño de flujos y componentes reutilizables, Design Systems.
7. NO OMITAS tecnologías del Master que la oferta valora: si la oferta pide o menciona un área (backend, IA, testing, cloud...) y el Master tiene una tecnología concreta de esa área, esa tecnología DEBE aparecer en HABILIDADES TÉCNICAS y, si encaja, en un bullet. Ejemplo: oferta backend/IA con Python y el Master incluye FastAPI → FastAPI debe salir en Backend. La REGLA DE EVIDENCIA impide inventar; esta regla impide lo contrario, dejarse fuera algo real y relevante.
8. Máximo 2 páginas

IDENTIDAD vs POSICIONAMIENTO (distinción base: no las confundas nunca):
- IDENTIDAD = quién ES la candidata. Es un repertorio CERRADO, declarado en "Identidades permitidas" del PERFIL BASE. No se amplía, no se deduce, no se negocia. Ejemplo de identidades: Frontend Tech Lead, Full-Stack Developer, AI Engineer.
- POSICIONAMIENTO = cómo se PRESENTA esa misma trayectoria ante esta oferta concreta. Es variable y lo fija el ARQUETIPO de la oferta. Ejemplos de posicionamiento: GenAI Adoption, Context Engineering, Applied AI, AI Automation, Design Systems.
- Un posicionamiento NO es una identidad nueva: es la MISMA trayectoria profesional presentada según el problema que la empresa quiere resolver. Por eso el posicionamiento puede cambiar en cada oferta y la identidad no cambia nunca.
- DÓNDE VA CADA COSA EN EL TITULAR: los huecos de IDENTIDAD solo admiten valores de "Identidades permitidas". Los huecos de MODIFICADOR (especialización y stack) son donde vive el posicionamiento. Un posicionamiento JAMÁS ocupa un hueco de identidad ni se coloca delante de las identidades.
- EL POSICIONAMIENTO TAMBIÉN NECESITA RESPALDO: un posicionamiento solo puede usarse si la EXPERIENCIA, proyectos o formación del Master lo respaldan. Un posicionamiento sin evidencia es una identidad inventada con otro nombre, y está igual de prohibido. Si el Master no respalda el posicionamiento que pide la oferta, se usa el que sí esté respaldado, aunque encaje peor.

HEADLINE RULES (primera línea del output — el TITULAR del CV es DATA-DRIVEN):
- EL PERFIL BASE ES UN CONTRATO, NO UNA SUGERENCIA: el bloque "PERFIL BASE" del CV MASTER declara la identidad de la candidata en secciones explícitas. Tu trabajo es LEERLAS, no interpretarlas. No deduzcas nada que esté declarado. Secciones del contrato:
  · "Identidad profesional" → el TITULAR BASE completo. Es el ancla.
  · "Identidades permitidas" → el repertorio CERRADO de identidades. Ninguna otra existe.
  · "Orden del titular" → el orden exacto de los elementos del titular. Es un dato, no una decisión tuya.
  · "Variante permitida" → el único titular alternativo, con la condición que lo habilita.
  · "Nunca permitido" → restricciones que el propio Master declara. Son innegociables.
- CÓMO CONSTRUIR EL TITULAR: parte de "Identidad profesional" tal cual está escrita y respeta el "Orden del titular". Tus ÚNICAS libertades son: (a) SUSTITUIR uno, como mucho dos, MODIFICADORES de especialización o stack por los que esta oferta valora, tomados siempre del Master; (b) OMITIR un modificador que no aporte nada a esta oferta. Las identidades y su orden no se tocan.
- EL ORDEN NO SE ALTERA: el orden de las identidades es branding, no una preferencia. Un recruiter que vea varios CV suyos debe reconocer el mismo titular. Prohibido reordenarlas aunque esta oferta priorice otra cosa, prohibido añadir identidades fuera de "Identidades permitidas", y prohibido reescribir el núcleo del titular para parecerse al título del anuncio.
- ÚNICA EXCEPCIÓN AL ORDEN: usa el titular de "Variante permitida" solo si esta oferta cumple EXACTAMENTE la condición que esa sección declara. Si el Master no declara variante, o la condición no se cumple, no hay excepción.
- LA OFERTA DECIDE QUÉ DESTACAR, NUNCA QUÉ INVENTAR: si la oferta pide un rol/identidad que NO está en "Identidades permitidas", NO lo uses. La oferta solo elige qué modificadores acompañan al titular.
- NADA DE ECO EN LAS IDENTIDADES: no copies calificativos ni adjetivos del título de la oferta a la identidad. Si la oferta se titula "Applied AI Engineer" y el PERFIL BASE dice "AI Engineer", el titular usa "AI Engineer", no "Applied AI Engineer". La identidad sale del PERFIL BASE tal cual está escrita ahí.
- GUARDRAIL POR PRINCIPIO (no escalar el NIVEL JERÁRQUICO): no incrementes el nivel jerárquico, la autoridad ni el alcance organizativo que declara el PERFIL BASE. La prueba NO es si la palabra aparece en una lista: es si el titular sugiere un rango, una autoridad o un alcance MAYOR que el declarado. Queda prohibido cualquier término que eleve el nivel, salvo que esa identidad EXACTA figure en el PERFIL BASE. Ejemplos, y la lista NO es cerrada: Principal, Staff, Head, Director, Architect, Distinguished, Manager, Leader, Chief, VP, Owner, Champion, Evangelist, Authority, o "Lead" de personas. Ante la duda, usa la identidad tal cual está escrita en el PERFIL BASE.
- RESPETA EL POSICIONAMIENTO: si el CV master incluye un bloque "POSICIONAMIENTO" (lo que la candidata NO es), el titular NUNCA debe contradecirlo.
- AUTO-CHEQUEO antes de cerrar el HEADLINE: verifica que CADA identidad del titular aparece en el PERFIL BASE; si alguna no rastrea ahí, elimínala.
- COHERENCIA IDENTIDAD/EXPERIENCIA: el PERFIL BASE define quién es la candidata; la sección EXPERIENCIA demuestra por qué puede afirmarlo. Cada identidad del titular debe poder justificarse leyendo la EXPERIENCIA del Master. Si una identidad del PERFIL BASE no tiene experiencia real que la respalde, NO la uses en el titular.
- FALLBACK: si el CV master NO contiene un bloque "PERFIL BASE", deriva las identidades SOLO de la experiencia real del CV master (nunca inventes) y aplica igualmente el guardrail de no escalar seniority.
- El titular va en el idioma de la oferta; separa identidades/skills con " | " o " · ".
- AÑOS DE EXPERIENCIA: usa el seniority tal como lo declara el PERFIL BASE (dentro de "Identidad profesional" o en una sección propia). No infles el número ni lo subas por encima de lo que dice la fuente.

RESUMEN — ESTABILIDAD (aproximadamente 70-80% estable, 20-30% adaptado): el resumen NO se reescribe desde cero en cada oferta. La mayor parte describe la MISMA trayectoria con las mismas ideas y casi las mismas palabras: de dónde viene la candidata, cómo ha evolucionado y qué la define hoy. Solo la parte final, o los ejemplos concretos que se eligen, se ajustan al arquetipo de esta oferta. El objetivo es que el titular, el resumen y su perfil público cuenten la misma historia; si alguien compara dos CV suyos, tiene que ver a la misma profesional enfocando distinto, nunca dos perfiles inconexos.

RESUMEN / PERFIL — no pierdas experiencia real que no cabe en el titular: si la candidata tiene fortalezas relevantes que el titular de esta oferta no refleja (según el CV Master), inclúyelas en el resumen para no perderlas, redactadas como experiencia real. Si el CV Master incluye un bloque "EVOLUCIÓN PROFESIONAL", úsalo para entender el arco de su carrera y dar el contexto temporal correcto (de dónde viene y hacia dónde ha evolucionado), sin inventar.

PERFIL — ANCLAJE A LA OFERTA (obligatorio): identifica 2-3 requisitos o palabras clave concretas de la DESCRIPCIÓN de la oferta que la candidata YA haya trabajado de verdad (según su CV master), e intégralos en el resumen redactados como experiencia real y demostrable ("con experiencia en X aplicada a Y", "habiendo trabajado Z en..."). PROHIBIDO incluir un requisito de la oferta que NO esté respaldado por su trayectoria real: si la oferta lo pide pero ella no lo ha hecho, NO entra. El objetivo es que el perfil resuene con la oferta usando SOLO lo que es cierto y defendible en entrevista.

PROYECTOS PROPIOS, FREELANCE Y CONSULTORÍA (no sobredimensionar la escala):
- Un proyecto personal, freelance o de consultoría se describe por la COMPLEJIDAD TÉCNICA del trabajo, NUNCA por el tamaño aparente de la organización. El lector debe entender QUÉ SABE HACER la candidata, no cómo de grande era la empresa.
- PROHIBIDO el lenguaje que sugiera equipos, departamentos o estructuras que no existían: "definí la estrategia de IA de la compañía", "lideré la arquitectura de la empresa", "responsable de la plataforma global", "dirección técnica de", "lideré un equipo de". Nada de vocabulario de CEO (estrategia, dirección, organización, transformación digital de la empresa) salvo que la oferta sea justamente para ese tipo de puesto.
- En su lugar, prioriza: qué construyó, qué problemas resolvió, qué tecnologías usó, qué arquitectura diseñó, qué decisiones de ingeniería tomó.
- EL RESUMEN NUNCA GIRA ALREDEDOR DEL PROYECTO PROPIO. El resumen describe la trayectoria COMPLETA; la experiencia actual es el EJEMPLO de la evolución, no el eje de la identidad. La narrativa correcta es un arco: años de trayectoria en el sector, especialización de origen, evolución posterior y especialización actual, todo ello leído del bloque "EVOLUCIÓN PROFESIONAL" o de la experiencia del Master. La narrativa INCORRECTA es identificar a la candidata con su proyecto más reciente ("fundadora de X que hace Y").
- El PESO de una experiencia no depende del tamaño de la empresa, sino de la RELEVANCIA de las competencias para esta oferta. Un proyecto propio puede ir el primero si es lo más reciente y especializado, pero presentado como trabajo de ingeniería, no como si hubiese dirigido una organización.

NIVEL DEL PUESTO (aplica al CUERPO del CV — el TITULAR lo fijan las HEADLINE RULES):
- Si el puesto NO menciona lead/manager/responsable/principal/head/coordinador/director → es un rol de DESARROLLO INDIVIDUAL. En ese caso, en el CUERPO:
  · REDUCE al mínimo el liderazgo: NO abras bullets con "Lideré/Coordiné equipos" ni "formación de equipos". Reformula esos logros hacia el trabajo TÉCNICO concreto (qué construiste, qué migraste, qué arquitectura/componentes/APIs), no hacia la gestión.
  · El liderazgo puede aparecer como contexto breve ("durante 8 años en el equipo frontend..."), NUNCA como la venta principal del perfil.
  · El titular puede reflejar seniority de liderazgo si figura en el PERFIL BASE; NO lo contradice, pero el cuerpo se centra en el trabajo técnico, no en dirigir personas.
- Solo si el puesto pide lead/manager/responsable/principal/head → destaca el ownership y la coordinación técnica.

ARQUETIPO DE LA OFERTA (decide QUÉ se proyecta; jamás QUÉ se inventa):
Antes de escribir, clasifica esta oferta en UN arquetipo leyendo el PUESTO y la DESCRIPCIÓN — no el sector de la empresa ni las tecnologías que nombra de pasada. El arquetipo NO cambia el titular base ni las identidades: cambia qué experiencia del Master va primero, qué bullets se priorizan y qué keywords entran.
- Frontend: React, Vue, TypeScript, arquitectura frontend, design systems, rendimiento, accesibilidad, mentoría técnica.
- Full Stack: el frontend como fortaleza principal, más Node, APIs, bases de datos e integración.
- Tech Lead: ownership técnico, estándares, code review, coordinación con producto, diseño y backend. No afirmes dirección de personas salvo que el Master lo respalde.
- UX Engineer: Figma, design systems, accesibilidad, colaboración con diseño.
- IA / AI Engineer: CONSTRUYE sistemas con IA. LLM, RAG, agentes, APIs, Context Engineering, evaluación, guardrails, pipelines.
- IA / GenAI Adoption: consigue que OTROS desarrolladores trabajen mejor con IA. Formación, workshops, mentoring, pairing, experimentación, herramientas de desarrollo asistido, playbooks, productividad de equipos de ingeniería.
- IA / AI Solutions Architect: DISEÑA sistemas. Arquitectura, escalabilidad, cloud, integración, decisiones técnicas, observabilidad, gobernanza.
- IA / AI Product Engineer: construye PRODUCTO con IA. Métricas, usuarios, experimentos, UX, negocio, iteración.
- IA / AI Automation Engineer: AUTOMATIZA procesos. N8N, MCP, APIs, workflows, integración de procesos de negocio.

"IA" NO es un arquetipo único. Una oferta que busca impulsar la adopción de GenAI en equipos de ingeniería y otra que busca construir sistemas LLM piden CV DISTINTOS aunque las dos digan "IA". Clasifica por el PROBLEMA que la empresa necesita resolver, no por la tecnología que menciona.

REGLA DE PROYECCIÓN: adapta el CV al PROBLEMA que resuelve la empresa que contrata, NO al producto que la candidata construyó. La misma trayectoria se proyecta hacia un arquetipo u otro sin inventar absolutamente nada: cambia el orden, el énfasis y qué se cuenta primero.

LÍMITE DEL ARQUETIPO: si el Master NO respalda el arquetipo de la oferta, no lo fuerces. Proyecta lo que haya y deja el resto fuera. Un arquetipo sin evidencia en el Master es una invitación a inventar, y esa línea no se cruza nunca: es preferible un CV honesto que encaja a medias, a uno que encaja del todo y no se sostiene en entrevista.

HECHOS, NO EFECTOS (separa lo que hizo de lo que eso demuestra):
- Escribe la ACCIÓN concreta y verificable, NUNCA el efecto que se le atribuye, salvo que el Master dé el dato. El lector deduce el efecto solo, y le convence más.
- MAL: "Improved engineering productivity", "Led AI transformation", "proven track record of measurable productivity gains", "measuring adoption impact", "drove AI adoption".
- BIEN: "Delivered technical workshops on Generative AI for engineering teams", "Designed and delivered internal training on prompt engineering".
- Prohibido el vocabulario de resultado no medido cuando el Master no trae el dato: "proven track record", "measurable", "impact", "transformation", "drove", "boosted", "accelerated". Un hecho concreto sin adjetivos vende más que un efecto declarado sin prueba, y además es defendible en entrevista.

POSICIONAMIENTO (adapta el ÉNFASIS a la oferta, nunca inventes):
- Prioriza y reordena las skills y logros del CV Master que esta oferta valora; deja en segundo plano lo que no pide. NUNCA añadas algo que no esté en el Master.
- NO MUEVAS SKILLS A EXPERIENCIA: una tecnología que el Master lista en HABILIDADES pero NO atribuye a un puesto concreto no puede aparecer como logro de ese puesto. Puede seguir en Habilidades, ahí es legítima. Atribuirla a una experiencia donde el Master no la sitúa es inventar, aunque la tecnología sea real y ella la domine.

- PERSONA GRAMATICAL (regla dura, aplica a TODO el CV, perfil y bullets): el CV se escribe SIN SUJETO, que es el estándar en CV anglosajón. Escribe "Frontend Tech Lead with 10+ years building...", "Led the migration from Vue to React", "Established engineering standards". PROHIBIDO en tercera persona ("She led", "Her career spans", "She brings") y PROHIBIDO en primera ("I led", "My career"). Esta instrucción está redactada hablando de la candidata en tercera persona por comodidad; el CV que produces NO debe heredar esa persona gramatical. Un CV que habla de ella desde fuera se lee como escrito por otro, y es de las cosas que más delatan que lo ha redactado una máquina.
- LÍMITE DE POSICIONAMIENTO: si el CV Master incluye un bloque "POSICIONAMIENTO" (lo que la candidata ES y lo que NO es), respétalo como frontera. NUNCA posiciones a la candidata en un rol o especialidad que ese bloque niega, aunque la oferta lo pida.
- EVOLUCIÓN: si el CV Master incluye un bloque "EVOLUCIÓN PROFESIONAL", respeta el arco temporal (de dónde viene, hacia dónde ha evolucionado); no presentes como especialidad actual algo que fue una etapa pasada, ni al revés.
- Una tecnología usada hace años puede presentarse como algo que puede retomar rápido por su experiencia previa, SIN presentarla como especialidad actual salvo que el Master lo respalde.

PASO 3 — REVISION ANTI-IA (aplicar al output antes de entregar):
Elimina TODO rastro de texto generado por IA:
- Cero guiones largos (—) ni dobles guiones (--)
- Cero frases tipo "responsable de...", "encargada de...", "orientada a..."
- Cero adjetivos vacíos ("dinámico", "proactivo", "apasionado", "motivado")
- Cero "passionate about", "I'd love to", "excited to"
- Cero verbos pasivos innecesarios ("fue responsable de..." → "lideró...")
- Si suena a IA, reescríbelo con lenguaje humano y directo
- Mantener tono profesional pero natural, como lo escribiría una persona

{bloque_formato}"""

    try:
        # Claude (calidad) primario; Groq de fallback dentro de call_llm_calidad
        respuesta_llm = call_llm_calidad(prompt, model=CV_MODEL, max_tokens=4096)
        contenido_cv  = respuesta_llm.contenido
    except RuntimeError as e:
        raise CVError(503, str(e))

    # 5. Limpiar output del LLM y extraer el titular (HEADLINE)
    #    Todo lo anterior a la línea HEADLINE se DESCARTA: si el modelo escribe el
    #    "ANÁLISIS INTERNO" pese al "no mostrar", o mete un preámbulo, jamás llega
    #    al DOCX. El cuerpo del CV empieza recién DESPUÉS del HEADLINE.
    titular = ""
    lineas = contenido_cv.split("\n")
    idx_headline = next(
        (i for i, l in enumerate(lineas)
         if l.strip().replace("**", "").replace("`", "").lower().startswith("headline:")),
        None,
    )
    if idx_headline is not None:
        cab = lineas[idx_headline].strip().replace("**", "").replace("`", "")
        titular = cab.split(":", 1)[1].strip()
        lineas = lineas[idx_headline + 1:]

    # 5a. El titular NO se acepta tal cual: se ENSAMBLA desde el PERFIL BASE. El
    #     contrato es cerrado (identidades y orden declarados) y tres formulaciones
    #     distintas de la regla en el prompt no lo sostuvieron: el modelo invirtio el
    #     orden, fusiono identidades con "&" y se comio la seniority. Del modelo solo
    #     se aprovechan los modificadores, y solo si el Master los respalda.
    titular = construir_titular(titular, cv_master)
    lineas_limpias = []
    for linea in lineas:
        limpia = linea.strip().replace("**", "").replace("`", "").replace("##", "").replace("# ", "")
        # Filtrar frases introductorias del LLM
        if limpia.lower().startswith(("aquí", "here is", "here's", "a continuación", "claro", "por supuesto")):
            continue
        lineas_limpias.append(limpia)
    contenido_cv = "\n".join(lineas_limpias)

    # 5b. Guardrail de veracidad: el prompt prohibe inventar metricas, pero eso es
    #     una peticion, no una garantia. Se contrasta la salida contra el Master.
    #     NO se aborta: una cifra sospechosa puede ser legitima (reformulacion) y
    #     abortar dejaria a la usuaria sin CV. Se avisa para que ella lo revise.
    cifras_sospechosas = detectar_cifras_no_respaldadas(contenido_cv, cv_master)
    if cifras_sospechosas:
        logger.warning(
            "CIFRAS NO RESPALDADAS por el CV Master en el CV de %s para %s/%s: %s",
            email, empresa, puesto, cifras_sospechosas,
        )

    #     Lo mismo con las tecnologias: la oferta pide una cosa y el modelo tiende a
    #     devolverla como si la candidata la tuviera. Tampoco se aborta, se avisa.
    tecnologias_sospechosas = detectar_tecnologias_no_respaldadas(contenido_cv, cv_master)
    if tecnologias_sospechosas:
        logger.warning(
            "TECNOLOGIAS NO RESPALDADAS por el CV Master en el CV de %s para %s/%s: %s",
            email, empresa, puesto, tecnologias_sospechosas,
        )

    #     Y la seccion de skills, elemento a elemento: es donde el modelo vuelca el
    #     stack de la oferta y donde el catalogo de arriba es ciego a lo que no
    #     tiene dado de alta, que es precisamente lo nuevo de cada oferta.
    skills_sospechosas = detectar_skills_no_respaldadas(contenido_cv, cv_master)
    if skills_sospechosas:
        logger.warning(
            "SKILLS NO RESPALDADAS por el CV Master en el CV de %s para %s/%s: %s",
            email, empresa, puesto, skills_sospechosas,
        )

    # 5d. Guardrail del titular: que respete el contrato del PERFIL BASE. Es el unico
    #     de los tres que mira la CABECERA, y viene de un fallo medido: los CV de N-iX y
    #     Revolut usaron la Variante permitida sin cumplir su condicion.
    titular_sospechoso = detectar_titular_fuera_de_contrato(titular, cv_master)
    if titular_sospechoso:
        logger.warning(
            "TITULAR FUERA DEL CONTRATO del PERFIL BASE en el CV de %s para %s/%s: %s",
            email, empresa, puesto, titular_sospechoso,
        )

    # 5e. Guardrail de ENTRADA (los otros tres miran la salida): ¿habia material que
    #     adaptar? Las ofertas de LinkedIn e Indeed llegan con 172-245 caracteres, el
    #     titular reformulado. Con eso el CV sale generico y ningun otro guardrail lo
    #     detecta, porque un CV generico no inventa nada: simplemente no dice nada.
    descripcion_evaluada = evaluar_descripcion_oferta(descripcion)
    if not descripcion_evaluada["suficiente"]:
        logger.warning(
            "DESCRIPCION INSUFICIENTE (%s chars) para %s/%s: %s",
            descripcion_evaluada["chars"], empresa, puesto, descripcion_evaluada["aviso"],
        )

    # 6. Generar DOCX con cabecera estructurada (titular adaptado por la oferta)
    nombre_archivo = _nombre_archivo_cv(nombre, puesto)
    docx_bytes = generar_docx_con_cabecera(contenido_cv, usuario, titular, idioma)

    # 7. Subir a Drive
    try:
        link_drive = subir_cv_a_drive(docx_bytes, nombre_archivo)
    except Exception as e:
        logger.error("Drive upload error: %s", e)
        raise CVError(500, f"Error subiendo a Drive: {e}")

    # 8. Anotar el enlace en la ficha de Notion, ahora, no via n8n. Si la cadena
    #    de n8n muere despues (timeout de 120s del nodo que llama aqui), el CV
    #    sigue estando localizable desde la ficha. Ver guardar_link_cv_en_notion.
    link_anotado = guardar_link_cv_en_notion(empresa, puesto, link_drive, nombre_archivo)

    return {
        "ok":              True,
        "link":            link_drive,
        # False = el CV existe en Drive pero no se pudo anotar en la ficha.
        # Revisar a mano: es un CV que Veronica tiene y no sabe que tiene.
        "link_anotado_en_notion": link_anotado,
        "modelo_usado":    respuesta_llm.modelo,
        "archivo":         nombre_archivo,
        "email":           email,
        "cv_master_usado": bool(cv_master),
        "idioma":          idioma,
        "cv_master_url":   master.url,
        # Vacio = todas las cifras del CV estan en el Master. Si trae algo, REVISAR
        # a mano antes de enviar: son datos que el modelo no pudo justificar.
        "cifras_no_respaldadas": cifras_sospechosas,
        # Vacio = todas las tecnologias del CV estan en el Master. Si trae algo, es una
        # tecnologia que la oferta pedia y la candidata NO tiene: quitarla antes de enviar.
        "tecnologias_no_respaldadas": tecnologias_sospechosas,
        # Vacio = el titular respeta el contrato del PERFIL BASE. Si trae algo, el titular
        # cambio la identidad o su orden: revisarlo, es lo primero que lee un recruiter.
        "titular_fuera_de_contrato": titular_sospechoso,
        # suficiente=False significa que la DESCRIPCION no daba material para adaptar:
        # el CV es generico aunque no haya ningun otro aviso. Es lo PRIMERO que hay que
        # mirar, porque los demas guardrails no detectan un CV correcto pero vacio.
        "descripcion_oferta": descripcion_evaluada,
    }


@app.route("/generar-cv", methods=["POST"])
def generar_cv():
    """Ruta Flask fina: parsea el request y delega en generar_cv_core (ADR-001)."""
    datos = request.get_json(force=True)
    try:
        result = generar_cv_core(
            email=datos.get("email", ""),
            empresa=datos.get("empresa", ""),
            puesto=datos.get("puesto", ""),
            descripcion=datos.get("descripcion", ""),
            idioma_in=(datos.get("idioma") or "").strip().lower(),
        )
    except CVError as e:
        return jsonify({"ok": False, "error": e.message}), e.status
    return jsonify(result)


@app.route("/generar-carta", methods=["POST"])
def generar_carta():
    """Genera la carta de presentación con la experiencia real del CV master.
    Usa Claude Sonnet (calidad) — la carta va a la empresa."""
    datos = request.get_json(force=True)
    email       = datos.get("email", "")
    empresa     = datos.get("empresa", "")
    puesto      = datos.get("puesto", "")
    descripcion = datos.get("descripcion", "")

    if not email or not empresa or not puesto:
        return jsonify({"ok": False, "error": "email, empresa y puesto son requeridos"}), 400

    # Idioma autoritativo (igual que en /generar-cv) + persona de contacto para
    # el saludo. Prioridad idioma: body.idioma > Idioma de Notion > detección.
    idioma_in = (datos.get("idioma") or "").strip().lower()
    contacto = (datos.get("contacto") or datos.get("nombre_contacto") or "").strip()
    if not descripcion.strip() or not contacto or idioma_in not in ("en", "es"):
        oferta = buscar_oferta_en_notion(empresa, puesto)
        if oferta:
            descripcion = descripcion or oferta.get("descripcion", "")
            contacto = contacto or oferta.get("nombre_contacto", "").strip()
            idioma_in = idioma_in or (oferta.get("idioma") or "").strip().lower()

    usuario = buscar_usuario_por_email(email)
    if not usuario:
        return jsonify({"ok": False, "error": f"Usuario {email} no encontrado en Notion"}), 404

    nombre = usuario.get("nombre") or email.split("@")[0]

    # Resolver idioma y leer el CV master en ese idioma
    idioma = idioma_in if idioma_in in ("en", "es") else idioma_de_oferta(puesto, descripcion, empresa)
    tiene_master = _tiene_algun_master(usuario)
    try:
        master    = leer_cv_master_desde_drive(usuario, idioma)
        cv_master = master.texto
    except Exception as e:
        # Mismo guard que /generar-cv: get_drive_service() -> creds.refresh() puede
        # fallar (token caducado o revocado) ANTES del try interno de la lectura.
        # Sin esto sale un 500 HTML opaco y el fallo tarda en diagnosticarse: paso
        # el 24-jul-2026, con el GOOGLE_REFRESH_TOKEN caducado, y dejo la cadena de
        # n8n rota sin ningún mensaje util.
        logger.error("Drive auth/lectura falló en /generar-carta: %s", e)
        return jsonify({
            "ok": False,
            "error": f"No se pudo autenticar/leer el CV master en Drive: {e}",
        }), 502

    def _es_legible(t: str) -> bool:
        if not t:
            return False
        if t.lstrip().startswith("PK"):
            return False
        imprimibles = sum(1 for c in t if c.isprintable() or c in "\n\r\t")
        return imprimibles / len(t) >= 0.85

    if tiene_master and not _es_legible(cv_master):
        logger.error("CV master ILEGIBLE para carta de %s — abortando", email)
        return jsonify({
            "ok": False,
            "error": ("No se pudo leer tu CV master desde Drive. NO se generó la carta "
                      "para evitar inventar datos. Revisá el archivo y permisos en Drive."),
        }), 502

    contexto = (f"CV MASTER (usa SOLO esta experiencia real, NO inventes nada):\n{cv_master}"
                if cv_master else
                f"PERFIL: {nombre} — {usuario.get('rol','')} — {usuario.get('perfil','')}")

    # Saludo: si hay persona de contacto real, dirigir la carta a ella; si no,
    # saludo genérico. NUNCA inventar un nombre.
    if contacto:
        instr_saludo = (f'saludo dirigido a la persona de contacto ("A la atención de {contacto}," '
                        f'en español, "Dear {contacto}," en inglés) — usa EXACTAMENTE ese nombre, no lo inventes ni lo cambies')
    else:
        instr_saludo = 'saludo formal genérico ("Estimados/as," en español, "Dear Hiring Team," en inglés)'

    prompt = f"""Eres un experto en cartas de presentación para ofertas de trabajo.
Escribe una carta de presentación profesional para {nombre}.

{contexto}

OFERTA:
- Empresa: {empresa}
- Puesto: {puesto}
- Descripción: {descripcion or "No disponible"}

REGLAS:
- La oferta está en {"inglés" if idioma == "en" else "español"}. Escribe TODA la carta en ese idioma (saludo, cuerpo y despedida).
- Máximo 250 palabras.
- Usa SOLO experiencia real del CV master, y SOLO la relevante para este puesto; conecta esa experiencia con lo que pide la oferta. NO inventes, NO exageres y NO afirmes nada difícil de defender en entrevista (ni gestión de equipos ni arquitectura que no haya hecho).
- NIVEL: si el puesto NO menciona lead/manager/responsable/principal/head, es un rol de DESARROLLO INDIVIDUAL: NO uses la coordinación/liderazgo de equipos como argumento principal (nada de "experiencia coordinando equipos técnicos"). Enfoca el encaje TÉCNICO real (stacks, full-stack, APIs, mobile, capacidad de aprender rápido el stack de la oferta).
- Tono profesional, directo y humano. Cero frases vacías de IA: nada de "apasionada",
  "proactiva", "soluciones innovadoras", "emocionada de la oportunidad", "dinámica".
- Menciona logros o tecnologías concretas del CV que encajen con la oferta.
- Formato carta: {instr_saludo} ... cuerpo ... despedida formal ("Atentamente," / "Sincerely,") seguida de "{nombre}".
- Devuelve SOLO el texto de la carta, sin encabezados ni comentarios."""

    try:
        respuesta_llm = call_llm_calidad(prompt, model=CARTA_MODEL, max_tokens=1500)
        carta         = respuesta_llm.contenido
    except RuntimeError as e:
        return jsonify({"ok": False, "error": str(e)}), 503

    # Limpiar frases introductorias del LLM
    carta = carta.strip()
    for pref in ("aquí tienes", "aquí está", "here is", "here's", "claro", "por supuesto"):
        if carta.lower().startswith(pref):
            carta = carta.split("\n", 1)[-1].strip()
            break
    # Fuera guiones largos y flechas: la carta va a la empresa. Regla NO NEGOCIABLE.
    carta = sanear_tipografia(carta, idioma)

    # GUARDRAILS DE LA CARTA. Hasta el 18-ago-2026 la carta salia SIN NINGUNO: los
    # cuatro detectores existentes se aplicaban solo a `contenido_cv`. Y la carta es
    # lo primero que lee un humano; el CV lo abren despues, si les interesas.
    # Igual que en /generar-cv, esto AVISA y no aborta: un aviso puede ser una
    # reformulacion legitima, y abortar dejaria a la usuaria sin carta.
    # `detectar_skills_no_respaldadas` queda fuera a proposito: lee lineas de skills
    # separadas por puntos, y una carta es prosa. Aplicarlo aqui daria ruido.
    avisos = []
    for nombre_regla, sospechosas in (
        ("EXPERIENCIA MAL ATRIBUIDA", detectar_experiencia_mal_atribuida(carta, cv_master)),
        ("TECNOLOGIAS NO RESPALDADAS", detectar_tecnologias_no_respaldadas(carta, cv_master)),
        ("CIFRAS NO RESPALDADAS", detectar_cifras_no_respaldadas(carta, cv_master)),
    ):
        if sospechosas:
            logger.warning(
                "%s en la CARTA de %s para %s/%s: %s",
                nombre_regla, email, empresa, puesto, sospechosas,
            )
            avisos.append({"regla": nombre_regla, "hallazgos": sospechosas})

    return jsonify({
        "ok":              True,
        "carta":           carta,
        "modelo_usado":    respuesta_llm.modelo,
        "email":           email,
        "cv_master_usado": bool(cv_master),
        "avisos":          avisos,
    })


@app.route("/usuarios", methods=["GET"])
def usuarios():
    """Consulta usuarios activos en Notion."""
    if not NOTION_DB_USUARIOS:
        return jsonify({"ok": False, "error": "NOTION_DB_USUARIOS no configurada"}), 500
    try:
        resp = requests.post(
            f"https://api.notion.com/v1/databases/{NOTION_DB_USUARIOS}/query",
            headers=notion_headers(),
            json={"filter": {"property": "Activo", "checkbox": {"equals": True}}},
            timeout=15,
        )
        resp.raise_for_status()
        results = resp.json().get("results", [])
        usuarios_list = []
        for p in results:
            props = p.get("properties", {})
            usuarios_list.append({
                "id":     p["id"],
                "nombre": props.get("Name", {}).get("title", [{}])[0].get("plain_text", ""),
                "email":  props.get("Email", {}).get("email", ""),
                "activo": props.get("Activo", {}).get("checkbox", False),
            })
        return jsonify({"ok": True, "usuarios": usuarios_list, "total": len(usuarios_list)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/buscar-ofertas-reales", methods=["POST"])
def buscar_ofertas_reales_endpoint():
    """
    Busca ofertas REALES en Remotive (sustituye al LLM inventando ofertas).

    Body esperado:
    {
        "rol": "frontend developer",        // opcional, default del perfil
        "stack": ["react", "typescript"],    // opcional
        "modalidad": ["Remoto"],             // opcional
        "ciudad": "Madrid",                  // opcional
    }
    """
    datos = request.get_json(force=True)
    perfil = datos.get("perfil", "")
    rol = datos.get("rol", "")
    stack = datos.get("stack", [])
    modalidad = datos.get("modalidad", [])
    ciudad = datos.get("ciudad", "")
    salario_min = datos.get("salario_min", 0)
    top_n = datos.get("top_n", 5)

    try:
        ofertas = buscar_ofertas_reales(perfil=perfil, rol=rol, stack=stack, salario_min=salario_min, modalidad=modalidad, ciudad=ciudad, top_n=top_n)
        return jsonify({"ok": True, "ofertas": ofertas, "total": len(ofertas)})
    except Exception as e:
        logger.error("Error buscando ofertas reales: %s", e)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/crear-oferta", methods=["POST"])
def crear_oferta():
    """Crea una oferta en Notion con TODOS sus campos.

    n8n llama aquí (una vez por oferta) en vez de mapear campo por campo en su
    nodo de Notion. Body: {"email": "...", "oferta": {...}} — o el objeto oferta
    plano. Detecta y persiste el idioma a partir de la descripción.
    """
    datos = request.get_json(force=True)
    email = (datos.get("email") or "").strip().lower()
    oferta = datos.get("oferta") if isinstance(datos.get("oferta"), dict) else datos

    empresa = oferta.get("empresa", "")
    puesto = oferta.get("puesto", "")
    if not empresa or not puesto:
        return jsonify({"ok": False, "error": "empresa y puesto son requeridos"}), 400

    # Detectar idioma ahora que tenemos la descripción en la mano y persistirlo
    idioma = idioma_de_oferta(puesto, oferta.get("descripcion", ""), empresa)

    # Relacionar la oferta con el usuario (si tenemos email)
    usuario_notion_id = ""
    if email:
        u = buscar_usuario_por_email(email)
        if u:
            usuario_notion_id = u.get("notion_id", "")

    try:
        page = crear_oferta_en_notion(oferta, idioma, usuario_notion_id)
    except requests.RequestException as e:
        logger.error("Error creando oferta en Notion: %s", e)
        return jsonify({"ok": False, "error": str(e)}), 502

    return jsonify({"ok": True, "notion_page_id": page.get("id", ""), "idioma": idioma})


# ══════════════════════════════════════════════
if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)