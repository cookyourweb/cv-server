"""Render del CV a DOCX, y el saneado tipografico que se aplica al escribirlo.

Convierte el texto que devuelve el LLM en un `.docx` con cabecera estructurada.
Entra texto, sale `bytes`: no sabe nada de Notion, de Drive ni del servidor.

`sanear_tipografia` vive aqui porque es una regla de RENDER: quita guiones largos
y flechas del texto que sale hacia la empresa, nunca del texto que el parser usa
para detectar la estructura del documento. Es una red determinista, no depende de
que el modelo obedezca el prompt.

Extraido de `cv_server_railway.py` el 28-ago-2026.
"""
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Guiones largos/medios y flechas: rastro tipográfico de IA. Regla NO NEGOCIABLE
# de la usuaria — jamás deben aparecer en un CV o carta que sale a una empresa.
_ARROWS = "→←⟶⟹➜➔➡⇒"
_DASHES = "—–―‒−"
_RE_ARROW = re.compile(r"\s*[" + _ARROWS + r"]\s*")
_RE_DASH  = re.compile(r"\s*[" + _DASHES + r"]\s*")
_RE_SPACES = re.compile(r"[ \t]{2,}")


def sanear_tipografia(texto: str, idioma: str = "es") -> str:
    """Elimina guiones largos/medios (—, –) y flechas (→) del texto final.

    Se aplica en el RENDER (DOCX y carta), nunca sobre el texto que el parser del
    DOCX usa para detectar estructura (ese sigue viendo el — crudo). Las flechas se
    traducen a la palabra de transición del idioma ("a"/"to"); los guiones a guion
    normal. Es una red determinista: no depende de que el LLM obedezca el prompt."""
    if not texto:
        return texto
    trans = " to " if idioma == "en" else " a "
    t = _RE_ARROW.sub(trans, texto)
    t = _RE_DASH.sub(" - ", t)
    t = _RE_SPACES.sub(" ", t)
    return t


def generar_docx(contenido_cv: str, nombre_candidato: str) -> bytes:
    """Wrapper legacy — usar generar_docx_con_cabecera() para nuevos CVs."""
    return generar_docx_con_cabecera(contenido_cv, {"nombre": nombre_candidato})


def generar_docx_con_cabecera(contenido_cv: str, usuario: dict, titular: str = "", idioma: str = "es") -> bytes:
    """Genera DOCX con cabecera profesional estructurada usando datos reales del usuario.
    `titular` (si viene) es el headline adaptado a la oferta por el LLM; tiene prioridad
    sobre el campo `rol` fijo del perfil."""
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BLUE = RGBColor(0x1A, 0x56, 0xDB)
    DARK = RGBColor(0x1A, 0x1A, 0x1A)
    GREY = RGBColor(0x66, 0x66, 0x66)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # ── Cabecera ──────────────────────────────────────────────────
    nombre   = usuario.get("nombre", "Candidato")
    rol      = titular or usuario.get("rol", "")
    ciudad   = usuario.get("ciudad", "")
    telefono = usuario.get("telefono", "")
    # Email de cabecera (contacto) separado del email-clave de búsqueda en Notion
    email    = usuario.get("email_cv") or usuario.get("email", "")
    linkedin = (usuario.get("linkedin", "") or "").replace("https://", "").replace("http://", "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(nombre.upper())
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = DARK

    if rol:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(sanear_tipografia(rol, idioma))
        r.font.size = Pt(11); r.font.color.rgb = BLUE

    contacto = " · ".join(c for c in [ciudad, telefono, email, linkedin] if c)
    if contacto:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(contacto)
        r.font.size = Pt(8.5); r.font.color.rgb = GREY

    # Línea separadora
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1A56DB")
    pBdr.append(bottom); pPr.append(pBdr)

    # ── Cuerpo del CV ────────────────────────────────────────────
    SECCIONES = ["PERFIL PROFESIONAL", "EXPERIENCIA PROFESIONAL", "EXPERIENCIA",
                 "HABILIDADES TÉCNICAS", "HABILIDADES", "FORMACIÓN", "IDIOMAS",
                 "PROYECTOS", "CERTIFICACIONES", "COMPETENCIAS"]

    for linea in contenido_cv.strip().split("\n"):
        linea = linea.strip()
        if not linea:
            continue

        limpia = linea.upper().strip()
        # La DETECCIÓN de estructura usa `linea` cruda (necesita ver el — como
        # marcador de empresa). El RENDER usa la versión saneada: así ningún
        # guion largo ni flecha llega nunca al DOCX. Regla NO NEGOCIABLE.
        render = sanear_tipografia(linea, idioma)

        # Sección
        if any(limpia.startswith(s) for s in SECCIONES) and len(linea) < 50:
            p = doc.add_paragraph()
            r = p.add_run(render.upper())
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLUE
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1A56DB")
            pBdr.append(bottom); pPr.append(pBdr)
            continue

        # Bullet
        if linea.startswith(("- ", "• ", "* ")):
            p = doc.add_paragraph()
            r = p.add_run("• " + render[2:].strip())
            r.font.size = Pt(9.5); r.font.color.rgb = DARK
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after  = Pt(2)
            continue

        # Empresa / puesto (línea con — o –)
        if ("—" in linea or "–" in linea) and len(linea) < 100:
            p = doc.add_paragraph()
            r = p.add_run(render)
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(1)
            continue

        # Fecha (línea corta con año)
        if re.search(r"(20\d{2}|19\d{2})", linea) and len(linea) < 60:
            p = doc.add_paragraph()
            r = p.add_run(render)
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY
            p.paragraph_format.space_after = Pt(2)
            continue

        # Texto normal
        p = doc.add_paragraph()
        r = p.add_run(render)
        r.font.size = Pt(9.5); r.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(3)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
